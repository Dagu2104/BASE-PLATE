import math
from dataclasses import dataclass, asdict

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import streamlit as st
from matplotlib.patches import Circle, Rectangle
import os
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

# ============================================================
# CONFIGURACIÓN GENERAL
# ============================================================

st.set_page_config(
    page_title="Base Plate Modular v1",
    layout="wide",
)

st.title("Diseño modular de placa base de acero")
st.caption("Versión modular en cadena: datos, geometría, módulo uniaxial preliminar y base para ampliación.")


# ============================================================
# DATACLASSES
# ============================================================

@dataclass
class Loads:
    Pu_kN: float
    Mux_kNm: float
    Muy_kNm: float
    Vux_kN: float = 0.0
    Vuy_kN: float = 0.0


@dataclass
class Materials:
    Fy_plate_MPa: float
    Fu_plate_MPa: float
    fc_MPa: float
    Fy_anchor_MPa: float
    Fu_anchor_MPa: float


@dataclass
class ColumnGeometry:
    d_col_mm: float
    bf_col_mm: float


@dataclass
class BasePlateGeometry:
    B_bp_mm: float
    N_bp_mm: float
    tp_mm: float


@dataclass
class AnchorLayout:
    nbx: int
    nby: int
    edge_x_mm: float
    edge_y_mm: float
    db_mm: float
    Ab_mm2: float
    hef_mm: float


@dataclass
class PedestalGeometry:
    B_ped_mm: float
    N_ped_mm: float
    h_ped_mm: float


# ============================================================
# FUNCIONES AUXILIARES GENERALES
# ============================================================

def validate_positive(value: float, name: str) -> None:
    if value <= 0:
        raise ValueError(f"'{name}' debe ser mayor que cero. Valor recibido: {value}")


def validate_nonnegative(value: float, name: str) -> None:
    if value < 0:
        raise ValueError(f"'{name}' no puede ser negativo. Valor recibido: {value}")


def validate_inputs(
    loads: Loads,
    materials: Materials,
    column: ColumnGeometry,
    base_plate: BasePlateGeometry,
    anchors: AnchorLayout,
    pedestal: PedestalGeometry,
) -> None:
    validate_positive(loads.Pu_kN, "Pu_kN")
    validate_nonnegative(abs(loads.Mux_kNm), "abs(Mux_kNm)")
    validate_nonnegative(abs(loads.Muy_kNm), "abs(Muy_kNm)")
    validate_nonnegative(abs(loads.Vux_kN), "abs(Vux_kN)")
    validate_nonnegative(abs(loads.Vuy_kN), "abs(Vuy_kN)")

    for name, value in asdict(materials).items():
        validate_positive(value, name)

    for name, value in asdict(column).items():
        validate_positive(value, name)

    for name, value in asdict(base_plate).items():
        validate_positive(value, name)

    if anchors.nbx < 2 or anchors.nby < 2:
        raise ValueError("Para disposición perimetral, nbx y nby deben ser al menos 2.")

    validate_positive(anchors.edge_x_mm, "edge_x_mm")
    validate_positive(anchors.edge_y_mm, "edge_y_mm")
    validate_positive(anchors.db_mm, "db_mm")
    validate_positive(anchors.Ab_mm2, "Ab_mm2")
    validate_positive(anchors.hef_mm, "hef_mm")

    for name, value in asdict(pedestal).items():
        validate_positive(value, name)

    if column.bf_col_mm > base_plate.B_bp_mm:
        raise ValueError(
            "El ancho de columna bf_col_mm no puede ser mayor que el ancho de placa B_bp_mm."
        )

    if column.d_col_mm > base_plate.N_bp_mm:
        raise ValueError(
            "El peralte de columna d_col_mm no puede ser mayor que el largo de placa N_bp_mm."
        )

    if base_plate.B_bp_mm > pedestal.B_ped_mm:
        raise ValueError(
            "La dimensión B de la placa no puede ser mayor que la dimensión B del pedestal."
        )

    if base_plate.N_bp_mm > pedestal.N_ped_mm:
        raise ValueError(
            "La dimensión N de la placa no puede ser mayor que la dimensión N del pedestal."
        )

    available_x = base_plate.B_bp_mm - 2.0 * anchors.edge_x_mm
    available_y = base_plate.N_bp_mm - 2.0 * anchors.edge_y_mm

    if available_x < 0:
        raise ValueError(
            "La distancia borde-eje de pernos en x es demasiado grande para la placa."
        )

    if available_y < 0:
        raise ValueError(
            "La distancia borde-eje de pernos en y es demasiado grande para la placa."
        )


def generate_bolt_coordinates(
    B_bp_mm: float,
    N_bp_mm: float,
    nbx: int,
    nby: int,
    edge_x_mm: float,
    edge_y_mm: float,
) -> pd.DataFrame:
    if nbx < 2 or nby < 2:
        raise ValueError("Para una disposición perimetral, nbx y nby deben ser al menos 2.")

    x_left = -B_bp_mm / 2.0 + edge_x_mm
    x_right = B_bp_mm / 2.0 - edge_x_mm
    y_bottom = -N_bp_mm / 2.0 + edge_y_mm
    y_top = N_bp_mm / 2.0 - edge_y_mm

    x_coords = np.linspace(x_left, x_right, nbx)
    y_coords = np.linspace(y_bottom, y_top, nby)

    bolts = []

    for x in x_coords:
        bolts.append((float(x), float(y_top)))

    for y in y_coords[-2:0:-1]:
        bolts.append((float(x_right), float(y)))

    for x in x_coords[::-1]:
        bolts.append((float(x), float(y_bottom)))

    for y in y_coords[1:-1]:
        bolts.append((float(x_left), float(y)))

    df = pd.DataFrame(bolts, columns=["x_mm", "y_mm"])
    df = df.drop_duplicates().reset_index(drop=True)
    df.insert(0, "Perno", np.arange(1, len(df) + 1))
    df["r_mm"] = np.sqrt(df["x_mm"] ** 2 + df["y_mm"] ** 2)

    return df


def build_input_summary(
    loads: Loads,
    materials: Materials,
    column: ColumnGeometry,
    base_plate: BasePlateGeometry,
    anchors: AnchorLayout,
    pedestal: PedestalGeometry,
) -> pd.DataFrame:
    rows = []

    for key, value in asdict(loads).items():
        rows.append(["Cargas", key, value])

    for key, value in asdict(materials).items():
        rows.append(["Materiales", key, value])

    for key, value in asdict(column).items():
        rows.append(["Columna", key, value])

    for key, value in asdict(base_plate).items():
        rows.append(["Placa base", key, value])

    for key, value in asdict(anchors).items():
        rows.append(["Pernos", key, value])

    for key, value in asdict(pedestal).items():
        rows.append(["Pedestal", key, value])

    return pd.DataFrame(rows, columns=["Bloque", "Variable", "Valor"])


def compute_layout_parameters(
    base_plate: BasePlateGeometry,
    anchors: AnchorLayout,
) -> dict:
    sx_mm = (base_plate.B_bp_mm - 2.0 * anchors.edge_x_mm) / (anchors.nbx - 1)
    sy_mm = (base_plate.N_bp_mm - 2.0 * anchors.edge_y_mm) / (anchors.nby - 1)
    total_bolts = 2 * anchors.nbx + 2 * anchors.nby - 4

    return {
        "A_plate_mm2": base_plate.B_bp_mm * base_plate.N_bp_mm,
        "total_bolts": total_bolts,
        "sx_mm": sx_mm,
        "sy_mm": sy_mm,
    }


# ============================================================
# MÓDULO 1 - ANÁLISIS UNIAXIAL PRELIMINAR
# ============================================================

def module1_uniaxial_preliminary(
    loads: Loads,
    base_plate: BasePlateGeometry,
    axis: str,
) -> dict:
    """
    Módulo 1:
    análisis preliminar uniaxial.

    Convención:
    - si axis = 'x', se usa Mux y la dimensión resistente N
      (flexión respecto al eje x produce variación de presión a lo largo de y)
    - si axis = 'y', se usa Muy y la dimensión resistente B
      (flexión respecto al eje y produce variación de presión a lo largo de x)
    """

    Pu_N = loads.Pu_kN * 1_000.0
    A_plate = base_plate.B_bp_mm * base_plate.N_bp_mm
    q_avg_MPa = Pu_N / A_plate

    if axis.lower() == "x":
        Mu_Nmm = loads.Mux_kNm * 1_000_000.0
        plate_dim_mm = base_plate.N_bp_mm
        axis_label = "x"
    elif axis.lower() == "y":
        Mu_Nmm = loads.Muy_kNm * 1_000_000.0
        plate_dim_mm = base_plate.B_bp_mm
        axis_label = "y"
    else:
        raise ValueError("El eje uniaxial debe ser 'x' o 'y'.")

    e_mm = Mu_Nmm / Pu_N
    kern_mm = plate_dim_mm / 6.0

    full_compression = abs(e_mm) <= kern_mm
    possible_uplift = abs(e_mm) > kern_mm

    return {
        "axis": axis_label,
        "Pu_kN": loads.Pu_kN,
        "Mu_kNm": Mu_Nmm / 1_000_000.0,
        "A_plate_mm2": A_plate,
        "q_avg_MPa": q_avg_MPa,
        "e_mm": e_mm,
        "kern_mm": kern_mm,
        "plate_dim_mm": plate_dim_mm,
        "full_compression": full_compression,
        "possible_uplift": possible_uplift,
        "e_over_kern": abs(e_mm) / kern_mm if kern_mm > 0 else float("nan"),
    }

# ============================================================
# MÓDULO 2 - COMPRESIÓN UNIAXIAL BAJO PLACA
# ============================================================

def module2_uniaxial_bearing(
    loads: Loads,
    base_plate: BasePlateGeometry,
    materials: Materials,
    axis: str,
) -> dict:
    """
    Módulo 2:
    análisis uniaxial de compresión bajo placa base.

    Convención:
    - axis = 'x'  -> se usa Mux y la variación de presión ocurre sobre N
    - axis = 'y'  -> se usa Muy y la variación de presión ocurre sobre B

    Hipótesis:
    - placa rígida
    - distribución lineal de presiones en compresión total
    - sin tracción del concreto
    - si e > L/6, se asume compresión parcial triangular

    Bearing conservador:
    q_allow_phi = phi * 0.85 * fc'
    con phi = 0.65
    """

    Pu_N = loads.Pu_kN * 1_000.0

    if axis.lower() == "x":
        Mu_Nmm = loads.Mux_kNm * 1_000_000.0
        L_mm = base_plate.N_bp_mm   # dimensión en la dirección de la flexión
        b_mm = base_plate.B_bp_mm   # ancho transversal
        axis_label = "x"
    elif axis.lower() == "y":
        Mu_Nmm = loads.Muy_kNm * 1_000_000.0
        L_mm = base_plate.B_bp_mm
        b_mm = base_plate.N_bp_mm
        axis_label = "y"
    else:
        raise ValueError("El eje uniaxial debe ser 'x' o 'y'.")

    A_mm2 = base_plate.B_bp_mm * base_plate.N_bp_mm
    e_mm = Mu_Nmm / Pu_N
    kern_mm = L_mm / 6.0

    phi_bearing = 0.65
    q_allow_phi_MPa = phi_bearing * 0.85 * materials.fc_MPa

    results = {
        "axis": axis_label,
        "Pu_kN": loads.Pu_kN,
        "Mu_kNm": Mu_Nmm / 1_000_000.0,
        "A_mm2": A_mm2,
        "L_mm": L_mm,
        "b_mm": b_mm,
        "e_mm": e_mm,
        "kern_mm": kern_mm,
        "phi_bearing": phi_bearing,
        "q_allow_phi_MPa": q_allow_phi_MPa,
    }

    # --------------------------------------------------------
    # CASO 1: COMPRESIÓN TOTAL
    # --------------------------------------------------------
    if abs(e_mm) <= kern_mm:
        q_avg_MPa = Pu_N / A_mm2

        # distribución lineal:
        # qmax = P/A * (1 + 6e/L)
        # qmin = P/A * (1 - 6e/L)
        factor = 6.0 * abs(e_mm) / L_mm
        q_max_MPa = q_avg_MPa * (1.0 + factor)
        q_min_MPa = q_avg_MPa * (1.0 - factor)

        # resultante de compresión coincide con la resultante aplicada
        # medida desde el centro de la placa
        xC_from_center_mm = e_mm

        results.update({
            "case": "full_compression",
            "q_avg_MPa": q_avg_MPa,
            "q_max_MPa": q_max_MPa,
            "q_min_MPa": q_min_MPa,
            "a_comp_mm": L_mm,
            "xC_from_center_mm": xC_from_center_mm,
            "bearing_ok": q_max_MPa <= q_allow_phi_MPa,
        })

    # --------------------------------------------------------
    # CASO 2: COMPRESIÓN PARCIAL
    # --------------------------------------------------------
    else:
        # longitud comprimida efectiva
        # a = 3*(L/2 - e) usando valor absoluto de e
        a_comp_mm = 3.0 * (L_mm / 2.0 - abs(e_mm))

        if a_comp_mm <= 0:
            raise ValueError(
                "La excentricidad es tan grande que la longitud comprimida efectiva resulta no positiva."
            )

        # presión máxima triangular:
        # qmax = 2P / (b * a)
        q_max_MPa = 2.0 * Pu_N / (b_mm * a_comp_mm)

        # qmin = 0 por levantamiento
        q_min_MPa = 0.0

        # ubicación de la resultante de compresión:
        # a/3 desde el borde comprimido
        # respecto al centro:
        toe_sign = 1.0 if e_mm >= 0 else -1.0
        xC_from_center_mm = toe_sign * (L_mm / 2.0 - a_comp_mm / 3.0)

        results.update({
            "case": "partial_compression",
            "q_avg_MPa": Pu_N / A_mm2,
            "q_max_MPa": q_max_MPa,
            "q_min_MPa": q_min_MPa,
            "a_comp_mm": a_comp_mm,
            "xC_from_center_mm": xC_from_center_mm,
            "bearing_ok": q_max_MPa <= q_allow_phi_MPa,
        })

    return results

# ============================================================
# MÓDULO 3 - TRACCIÓN UNIAXIAL EN PERNOS
# ============================================================

def module3_uniaxial_anchor_tension(
    loads: Loads,
    base_plate: BasePlateGeometry,
    anchors: AnchorLayout,
    bolt_df: pd.DataFrame,
    module2_results: dict,
) -> dict:
    """
    Módulo 3:
    cálculo preliminar de tracción en pernos para análisis uniaxial.

    Hipótesis:
    - si no hay levantamiento, T = 0
    - si hay compresión parcial, el momento lo equilibran:
        * compresión en concreto
        * tracción en la fila extrema de pernos del lado traccionado
    - la tracción total se reparte uniformemente entre los pernos
      de la fila extrema traccionada

    Convención:
    - axis = 'x' -> flexión asociada a Mux, variación sobre eje y
    - axis = 'y' -> flexión asociada a Muy, variación sobre eje x
    """

    axis = module2_results["axis"]
    Pu_N = loads.Pu_kN * 1_000.0

    if axis == "x":
        Mu_Nmm = loads.Mux_kNm * 1_000_000.0
        coord_name = "y_mm"
        extreme_coord = bolt_df["y_mm"].abs().max()
    elif axis == "y":
        Mu_Nmm = loads.Muy_kNm * 1_000_000.0
        coord_name = "x_mm"
        extreme_coord = bolt_df["x_mm"].abs().max()
    else:
        raise ValueError("El eje del Módulo 3 debe ser 'x' o 'y'.")

    # Si no hay compresión parcial, no se activa la tracción en este modelo
    if module2_results["case"] == "full_compression":
        return {
            "axis": axis,
            "case": "full_compression",
            "tension_active": False,
            "T_total_kN": 0.0,
            "n_tension_bolts": 0,
            "T_per_bolt_kN": 0.0,
            "lever_arm_mm": 0.0,
            "tension_side": "No aplica",
            "critical_bolts": [],
        }

    # --------------------------------------------------------
    # Lado traccionado
    # --------------------------------------------------------
    # Si e > 0, la compresión se va hacia el lado positivo.
    # La tracción aparece en el lado opuesto.
    e_mm = module2_results["e_mm"]

    if e_mm >= 0:
        tension_side = "negative"
        tension_coord = -extreme_coord
    else:
        tension_side = "positive"
        tension_coord = extreme_coord

    # Pernos activos en tracción = fila extrema del lado que levanta
    tol = 1e-6
    tension_bolts_df = bolt_df[np.isclose(bolt_df[coord_name], tension_coord, atol=tol)].copy()

    n_tension_bolts = len(tension_bolts_df)

    if n_tension_bolts == 0:
        raise ValueError("No se identificaron pernos en la fila extrema traccionada.")

    # --------------------------------------------------------
    # Brazo entre C y T
    # --------------------------------------------------------
    xC_from_center_mm = module2_results["xC_from_center_mm"]
    lever_arm_mm = abs(tension_coord - xC_from_center_mm)

    if lever_arm_mm <= 0:
        raise ValueError("El brazo interno entre compresión y tracción no es positivo.")

    # --------------------------------------------------------
    # Equilibrio de momentos
    # M = C*zC + T*zT
    # pero C = Pu + T  si tomamos equilibrio vertical con compresión positiva
    #
    # Para esta etapa preliminar, se usa equilibrio directo del par interno:
    # T_total = Mu / z
    #
    # Esto es una aproximación preliminar razonable para depuración
    # antes del refinamiento completo.
    # --------------------------------------------------------
    T_total_N = abs(Mu_Nmm) / lever_arm_mm
    T_per_bolt_N = T_total_N / n_tension_bolts

    critical_bolts = tension_bolts_df["Perno"].astype(int).tolist()

    return {
        "axis": axis,
        "case": "partial_compression",
        "tension_active": True,
        "T_total_kN": T_total_N / 1_000.0,
        "n_tension_bolts": n_tension_bolts,
        "T_per_bolt_kN": T_per_bolt_N / 1_000.0,
        "lever_arm_mm": lever_arm_mm,
        "tension_side": tension_side,
        "critical_bolts": critical_bolts,
        "tension_coord_mm": tension_coord,
    }

# ============================================================
# MÓDULO 4 - ESPESOR DE PLACA
# ============================================================

def module4_plate_thickness(
    base_plate: BasePlateGeometry,
    column: ColumnGeometry,
    materials: Materials,
    module2_results: dict,
) -> dict:
    """
    Módulo 4:
    cálculo preliminar del espesor mínimo de la placa base.

    Hipótesis:
    - la placa trabaja como franja en voladizo
    - se usa la presión máxima de contacto del Módulo 2 como presión gobernante
    - se toma el voladizo crítico entre ambas direcciones
    - resistencia a flexión simplificada de la placa:
          t_req = m * sqrt( 2*q / (phi*Fy) )

    Unidades:
    - q en MPa = N/mm²
    - Fy en MPa
    - m en mm
    - t_req en mm
    """

    B_mm = base_plate.B_bp_mm
    N_mm = base_plate.N_bp_mm
    tp_mm = base_plate.tp_mm

    bf_mm = column.bf_col_mm
    d_mm = column.d_col_mm

    Fy_MPa = materials.Fy_plate_MPa
    phi_flexure = 0.90

    # Voladizos geométricos libres
    mx_mm = (B_mm - bf_mm) / 2.0
    my_mm = (N_mm - d_mm) / 2.0

    if mx_mm <= 0:
        raise ValueError("El voladizo libre mx no es positivo. Revisa B y bf.")
    if my_mm <= 0:
        raise ValueError("El voladizo libre my no es positivo. Revisa N y d.")

    # Voladizo crítico
    mcrit_mm = max(mx_mm, my_mm)

    # Presión gobernante obtenida del módulo 2
    q_u_MPa = module2_results["q_max_MPa"]

    # Espesor mínimo requerido
    t_req_mm = mcrit_mm * math.sqrt((2.0 * q_u_MPa) / (phi_flexure * Fy_MPa))

    # Utilización
    thickness_ok = tp_mm >= t_req_mm
    utilization = t_req_mm / tp_mm if tp_mm > 0 else float("inf")

    return {
        "B_mm": B_mm,
        "N_mm": N_mm,
        "bf_mm": bf_mm,
        "d_mm": d_mm,
        "mx_mm": mx_mm,
        "my_mm": my_mm,
        "mcrit_mm": mcrit_mm,
        "q_u_MPa": q_u_MPa,
        "Fy_MPa": Fy_MPa,
        "phi_flexure": phi_flexure,
        "tp_input_mm": tp_mm,
        "t_req_mm": t_req_mm,
        "thickness_ok": thickness_ok,
        "utilization": utilization,
    }
# ============================================================
# MÓDULO 5 - ACERO DEL ANCLAJE
# ============================================================

def module5_anchor_steel_strength(
    loads: Loads,
    materials: Materials,
    anchors: AnchorLayout,
    layout_info: dict,
    module3_results: dict,
    analysis_mode: str,
    uniaxial_axis: str,
    anchor_type: str,
    phi_anchor_tension_steel: float,
    phi_anchor_shear_steel: float,
    use_built_up_grout_pad: bool,
) -> dict:
    """
    Módulo 5:
    revisión preliminar del acero del anclaje.

    Base:
    - Tensión: Nsa = Ase,N * futa
    - Cortante:
        headed_stud -> Vsa = Ase,V * futa
        headed_bolt / hooked_bolt / adhesive_anchor -> Vsa = 0.6 * Ase,V * futa
    - Si built-up grout pad: Vsa *= 0.80
    - Interacción acero: (Nua / phiNsa)^2 + (Vua / phiVsa)^2 <= 1

    Hipótesis actuales:
    - se usa Ab_mm2 como área efectiva tanto para tensión como para cortante
      en esta etapa preliminar
    - en uniaxial, el cortante demandado se toma como:
        axis = x -> Vux
        axis = y -> Vuy
    - el cortante se reparte uniformemente entre todos los pernos
    - la tensión se toma del perno crítico del Módulo 3
    """

    if analysis_mode != "Uniaxial":
        return None

    # --------------------------------------------------------
    # Área efectiva preliminar
    # --------------------------------------------------------
    AseN_mm2 = anchors.Ab_mm2
    AseV_mm2 = anchors.Ab_mm2

    # --------------------------------------------------------
    # Resistencia efectiva del material del anclaje
    # futa <= min(Fu, 1.9*Fy, 860 MPa)
    # --------------------------------------------------------
    futa_eff_MPa = min(
        materials.Fu_anchor_MPa,
        1.9 * materials.Fy_anchor_MPa,
        860.0
    )

    # --------------------------------------------------------
    # Resistencia nominal en tensión
    # --------------------------------------------------------
    Nsa_N = AseN_mm2 * futa_eff_MPa

    # --------------------------------------------------------
    # Resistencia nominal en cortante
    # --------------------------------------------------------
    if anchor_type == "headed_stud":
        Vsa_N = AseV_mm2 * futa_eff_MPa
    elif anchor_type in ["headed_bolt", "hooked_bolt", "adhesive_anchor"]:
        Vsa_N = 0.60 * AseV_mm2 * futa_eff_MPa
    else:
        raise ValueError("Tipo de anclaje no reconocido en Módulo 5.")

    if use_built_up_grout_pad:
        Vsa_N *= 0.80

    # --------------------------------------------------------
    # Resistencias de diseño
    # --------------------------------------------------------
    phiNsa_N = phi_anchor_tension_steel * Nsa_N
    phiVsa_N = phi_anchor_shear_steel * Vsa_N

    # --------------------------------------------------------
    # Demanda de tensión por perno crítico
    # --------------------------------------------------------
    if module3_results["tension_active"]:
        Nua_per_bolt_N = module3_results["T_per_bolt_kN"] * 1_000.0
    else:
        Nua_per_bolt_N = 0.0

    # --------------------------------------------------------
    # Demanda de cortante por perno
    # --------------------------------------------------------
    total_bolts = layout_info["total_bolts"]

    if uniaxial_axis == "x":
        Vua_total_N = abs(loads.Vux_kN) * 1_000.0
    elif uniaxial_axis == "y":
        Vua_total_N = abs(loads.Vuy_kN) * 1_000.0
    else:
        raise ValueError("El eje uniaxial debe ser 'x' o 'y'.")

    Vua_per_bolt_N = Vua_total_N / total_bolts if total_bolts > 0 else 0.0

    # --------------------------------------------------------
    # Interacción acero-acero
    # --------------------------------------------------------
    tension_ratio = Nua_per_bolt_N / phiNsa_N if phiNsa_N > 0 else float("inf")
    shear_ratio = Vua_per_bolt_N / phiVsa_N if phiVsa_N > 0 else float("inf")

    interaction_value = tension_ratio**2 + shear_ratio**2

    tension_ok = Nua_per_bolt_N <= phiNsa_N
    shear_ok = Vua_per_bolt_N <= phiVsa_N
    interaction_ok = interaction_value <= 1.0

    return {
        "anchor_type": anchor_type,
        "AseN_mm2": AseN_mm2,
        "AseV_mm2": AseV_mm2,
        "fya_MPa": materials.Fy_anchor_MPa,
        "futa_input_MPa": materials.Fu_anchor_MPa,
        "futa_eff_MPa": futa_eff_MPa,
        "phi_anchor_tension_steel": phi_anchor_tension_steel,
        "phi_anchor_shear_steel": phi_anchor_shear_steel,
        "Nsa_kN": Nsa_N / 1_000.0,
        "Vsa_kN": Vsa_N / 1_000.0,
        "phiNsa_kN": phiNsa_N / 1_000.0,
        "phiVsa_kN": phiVsa_N / 1_000.0,
        "Nua_per_bolt_kN": Nua_per_bolt_N / 1_000.0,
        "Vua_per_bolt_kN": Vua_per_bolt_N / 1_000.0,
        "tension_ratio": tension_ratio,
        "shear_ratio": shear_ratio,
        "interaction_value": interaction_value,
        "tension_ok": tension_ok,
        "shear_ok": shear_ok,
        "interaction_ok": interaction_ok,
        "use_built_up_grout_pad": use_built_up_grout_pad,
        "total_bolts": total_bolts,
    }

# ============================================================
# MÓDULO 6 - CONCRETO EN TENSIÓN (ACI 318-25)
# ============================================================

def module6_concrete_tension(
    materials: Materials,
    anchors: AnchorLayout,
    pedestal: PedestalGeometry,
    bolt_df: pd.DataFrame,
    module3_results: dict,
    anchor_type: str,
    anchor_installation: str,
    service_cracked: bool,
    lambda_a: float,
    psi_a: float,
    phi_concrete_tension: float,
    Abrg_mm2: float,
    eh_mm: float,
    ca1_x_mm: float,
    ca1_y_mm: float,
    ca2_mm: float,
) -> dict:
    """
    Módulo 6:
    revisión de anclaje al concreto en tensión.

    Alcance actual:
    - concrete breakout en tensión
    - pullout
    - side-face blowout
    - interacción de comparación contra la demanda de tensión del Módulo 3

    Validez práctica actual:
    - módulo pensado primero para grupos sin borde cercano crítico
    - no incluye todavía todos los modificadores completos de 17.6.2.3 a 17.6.2.5
    - no incluye bond de adhesive anchors
    - no incluye aún cortante de concreto ni pryout
    """

    # --------------------------------------------------------
    # Demanda de tensión del grupo y del perno crítico
    # --------------------------------------------------------
    if not module3_results["tension_active"]:
        Nua_per_bolt_N = 0.0
        Nua_group_N = 0.0
    else:
        Nua_per_bolt_N = module3_results["T_per_bolt_kN"] * 1_000.0
        Nua_group_N = module3_results["T_total_kN"] * 1_000.0

    hef_mm = anchors.hef_mm
    da_mm = anchors.db_mm
    fc_MPa = materials.fc_MPa

    # --------------------------------------------------------
    # k_c para concrete breakout
    # --------------------------------------------------------
    if anchor_installation == "cast_in":
        kc = 10.0
    elif anchor_installation == "post_installed":
        kc = 7.0
    else:
        raise ValueError("anchor_installation debe ser 'cast_in' o 'post_installed'.")

    # --------------------------------------------------------
    # A_Nco = 9 hef^2
    # --------------------------------------------------------
    ANco_mm2 = 9.0 * hef_mm**2

    # --------------------------------------------------------
    # A_Nc aproximado del grupo sin borde crítico
    # proyección simplificada rectangular:
    # (rango x del grupo + 3hef) * (rango y del grupo + 3hef)
    # esto equivale a extender 1.5hef por cada lado
    # --------------------------------------------------------
    x_min = bolt_df["x_mm"].min()
    x_max = bolt_df["x_mm"].max()
    y_min = bolt_df["y_mm"].min()
    y_max = bolt_df["y_mm"].max()

    span_x = x_max - x_min
    span_y = y_max - y_min

    ANc_mm2 = (span_x + 3.0 * hef_mm) * (span_y + 3.0 * hef_mm)

    # --------------------------------------------------------
    # Verificación simple de borde para uso del módulo
    # Si el borde disponible es menor que 1.5hef, advertimos que
    # el módulo está fuera de su rango ideal actual.
    # --------------------------------------------------------
    edge_ok_x = ca1_x_mm >= 1.5 * hef_mm
    edge_ok_y = ca1_y_mm >= 1.5 * hef_mm
    no_edge_effect_assumption_ok = edge_ok_x and edge_ok_y

    # --------------------------------------------------------
    # Concrete breakout básico
    # Nb = kc * lambda_a * sqrt(fc') * hef^1.5
    # --------------------------------------------------------
    Nb_N = kc * lambda_a * math.sqrt(fc_MPa) * (hef_mm ** 1.5)

    # Como versión inicial:
    # psi_ec,N = 1.0
    # psi_ed,N = 1.0
    # psi_c,N  = 1.0 (conservador para región agrietada)
    # psi_cp,N = 1.0
    psi_ec_N = 1.0
    psi_ed_N = 1.0
    psi_cp_N = 1.0
    psi_c_N = 1.0

    Ncbg_N = (ANc_mm2 / ANco_mm2) * psi_ec_N * psi_ed_N * psi_c_N * psi_cp_N * Nb_N
    phiNcbg_N = phi_concrete_tension * Ncbg_N

    # --------------------------------------------------------
    # Pullout
    # --------------------------------------------------------
    psi_c_P = 1.0 if service_cracked else 1.4

    Npn_N = None

    if anchor_type in ["headed_bolt", "headed_stud"] and anchor_installation == "cast_in":
        Np_N = 8.0 * Abrg_mm2 * fc_MPa
        Npn_N = psi_c_P * Np_N

    elif anchor_type == "hooked_bolt" and anchor_installation == "cast_in":
        if eh_mm < 3.0 * da_mm or eh_mm > 4.5 * da_mm:
            raise ValueError(
                f"Para hooked bolts, ACI 17.6.3.2.2(b) requiere 3da ≤ eh ≤ 4.5da. "
                f"Actualmente: eh={eh_mm:.3f} mm, da={da_mm:.3f} mm."
            )
        Np_N = 0.9 * fc_MPa * eh_mm * da_mm
        Npn_N = psi_c_P * Np_N

    elif anchor_type == "adhesive_anchor":
        # bond strength no se puede cerrar aquí sin datos del sistema/calificación
        Npn_N = None

    else:
        # Para post-installed expansion/screw/undercut, ACI exige valores por ensayo/calificación
        Npn_N = None

    phiNpn_N = phi_concrete_tension * Npn_N if Npn_N is not None else None

    # --------------------------------------------------------
    # Side-face blowout
    # Solo headed anchors
    # --------------------------------------------------------
    Nsbg_N = None
    phiNsbg_N = None

    if anchor_type in ["headed_bolt", "headed_stud"]:
        # Tomamos ca1 mínimo como el más desfavorable entre x e y
        ca1_mm = min(ca1_x_mm, ca1_y_mm)

        Nsb_single_N = psi_a * 13.0 * ca1_mm * math.sqrt(Abrg_mm2) * lambda_a * math.sqrt(fc_MPa)

        if ca2_mm < 3.0 * ca1_mm:
            ratio = ca2_mm / ca1_mm
            if ratio < 1.0:
                ratio = 1.0
            Nsb_single_N *= (1.0 + ratio) / 4.0

        # grupo: aproximación lineal conservadora con anclajes en tracción
        if module3_results["tension_active"]:
            n_tension = module3_results["n_tension_bolts"]
        else:
            n_tension = 0

        Nsbg_N = Nsb_single_N * max(n_tension, 1)
        phiNsbg_N = phi_concrete_tension * Nsbg_N

    # --------------------------------------------------------
    # Resistencia concreta gobernante en tensión
    # Para este módulo:
    # min(concrete breakout, pullout si aplica, side-face blowout si aplica)
    # --------------------------------------------------------
    resistance_candidates = [phiNcbg_N]

    if phiNpn_N is not None:
        resistance_candidates.append(phiNpn_N)

    if phiNsbg_N is not None:
        resistance_candidates.append(phiNsbg_N)

    phiNn_cg_N = min(resistance_candidates)

    concrete_tension_ok = Nua_group_N <= phiNn_cg_N

    return {
        "kc": kc,
        "lambda_a": lambda_a,
        "psi_a": psi_a,
        "phi_concrete_tension": phi_concrete_tension,
        "hef_mm": hef_mm,
        "da_mm": da_mm,
        "fc_MPa": fc_MPa,
        "ANco_mm2": ANco_mm2,
        "ANc_mm2": ANc_mm2,
        "edge_ok_x": edge_ok_x,
        "edge_ok_y": edge_ok_y,
        "no_edge_effect_assumption_ok": no_edge_effect_assumption_ok,
        "Nb_kN": Nb_N / 1_000.0,
        "Ncbg_kN": Ncbg_N / 1_000.0,
        "phiNcbg_kN": phiNcbg_N / 1_000.0,
        "psi_c_P": psi_c_P,
        "Npn_kN": None if Npn_N is None else Npn_N / 1_000.0,
        "phiNpn_kN": None if phiNpn_N is None else phiNpn_N / 1_000.0,
        "Nsbg_kN": None if Nsbg_N is None else Nsbg_N / 1_000.0,
        "phiNsbg_kN": None if phiNsbg_N is None else phiNsbg_N / 1_000.0,
        "Nua_per_bolt_kN": Nua_per_bolt_N / 1_000.0,
        "Nua_group_kN": Nua_group_N / 1_000.0,
        "phiNn_cg_kN": phiNn_cg_N / 1_000.0,
        "concrete_tension_ok": concrete_tension_ok,
        "anchor_type": anchor_type,
        "anchor_installation": anchor_installation,
    }

# ============================================================
# MÓDULO 7 - CONCRETO EN CORTANTE (ACI 17 PARCIAL)
# ============================================================

def module7_concrete_shear(
    loads: Loads,
    anchors: AnchorLayout,
    bolt_df: pd.DataFrame,
    module6_results: dict,
    analysis_mode: str,
    uniaxial_axis: str,
    lambda_a: float,
    phi_concrete_shear: float,
    ca1_shear_mm: float,
    member_thickness_for_shear_mm: float,
) -> dict:
    """
    Módulo 7:
    revisión preliminar del concreto en cortante.

    Incluye:
    - concrete breakout en shear (forma simplificada provisional)
    - pryout
    - interacción concreto tensión-cortante con exponente 5/3

    Importante:
    Esta versión todavía NO incluye todos los modificadores completos
    de ACI 17.7.2; se usa una formulación simplificada para depuración.
    """

    if analysis_mode != "Uniaxial":
        return None

    fc_MPa = module6_results["fc_MPa"]
    hef_mm = module6_results["hef_mm"]

    # --------------------------------------------------------
    # Cortante demandado del grupo
    # --------------------------------------------------------
    if uniaxial_axis == "x":
        Vua_group_N = abs(loads.Vux_kN) * 1_000.0
    elif uniaxial_axis == "y":
        Vua_group_N = abs(loads.Vuy_kN) * 1_000.0
    else:
        raise ValueError("El eje uniaxial debe ser 'x' o 'y'.")

    # --------------------------------------------------------
    # Área proyectada simplificada para shear breakout
    # Tomamos la extensión del grupo en dirección perpendicular al cortante
    # más 1.5 ca1 a cada lado.
    # --------------------------------------------------------
    if uniaxial_axis == "x":
        # cortante en x -> extensión perpendicular ~ y
        y_min = bolt_df["y_mm"].min()
        y_max = bolt_df["y_mm"].max()
        group_span_perp_mm = y_max - y_min
    else:
        # cortante en y -> extensión perpendicular ~ x
        x_min = bolt_df["x_mm"].min()
        x_max = bolt_df["x_mm"].max()
        group_span_perp_mm = x_max - x_min

    # área proyectada simplificada
    AVco_mm2 = 4.5 * ca1_shear_mm**2
    AVc_mm2 = (group_span_perp_mm + 3.0 * ca1_shear_mm) * (1.5 * ca1_shear_mm)

    # --------------------------------------------------------
    # Concrete breakout básico en shear
    # FORMA SIMPLIFICADA PROVISIONAL
    # --------------------------------------------------------
    Vb_N = 16.0 * lambda_a * math.sqrt(fc_MPa) * (ca1_shear_mm ** 1.5)

    Vcbg_N = (AVc_mm2 / AVco_mm2) * Vb_N
    phiVcbg_N = phi_concrete_shear * Vcbg_N

    # --------------------------------------------------------
    # Pryout
    # ACI commentary: kcp ~ 1 a 2
    # menor valor para anclajes cortos
    # --------------------------------------------------------
    if hef_mm < 63.5:
        kcp = 1.0
    else:
        kcp = 2.0

    # Se usa Ncbg del módulo 6 para estimar pryout
    Vcpg_N = kcp * (module6_results["Ncbg_kN"] * 1_000.0)
    phiVcpg_N = phi_concrete_shear * Vcpg_N

    # --------------------------------------------------------
    # Resistencia gobernante en cortante del concreto
    # --------------------------------------------------------
    phiVn_cg_N = min(phiVcbg_N, phiVcpg_N)

    shear_concrete_ok = Vua_group_N <= phiVn_cg_N

    # --------------------------------------------------------
    # Interacción concreto N-V con exponente 5/3
    # --------------------------------------------------------
    phiNn_cg_N = module6_results["phiNn_cg_kN"] * 1_000.0
    Nua_group_N = module6_results["Nua_group_kN"] * 1_000.0

    tension_term = (Nua_group_N / phiNn_cg_N) ** (5.0 / 3.0) if phiNn_cg_N > 0 else float("inf")
    shear_term = (Vua_group_N / phiVn_cg_N) ** (5.0 / 3.0) if phiVn_cg_N > 0 else float("inf")

    interaction_concrete = tension_term + shear_term
    interaction_concrete_ok = interaction_concrete <= 1.0

    # --------------------------------------------------------
    # Advertencia de espesor de miembro
    # --------------------------------------------------------
    thickness_warning = member_thickness_for_shear_mm < 1.5 * ca1_shear_mm

    return {
        "phi_concrete_shear": phi_concrete_shear,
        "fc_MPa": fc_MPa,
        "hef_mm": hef_mm,
        "lambda_a": lambda_a,
        "ca1_shear_mm": ca1_shear_mm,
        "AVco_mm2": AVco_mm2,
        "AVc_mm2": AVc_mm2,
        "Vb_kN": Vb_N / 1_000.0,
        "Vcbg_kN": Vcbg_N / 1_000.0,
        "phiVcbg_kN": phiVcbg_N / 1_000.0,
        "kcp": kcp,
        "Vcpg_kN": Vcpg_N / 1_000.0,
        "phiVcpg_kN": phiVcpg_N / 1_000.0,
        "phiVn_cg_kN": phiVn_cg_N / 1_000.0,
        "Vua_group_kN": Vua_group_N / 1_000.0,
        "shear_concrete_ok": shear_concrete_ok,
        "Nua_group_kN": Nua_group_N / 1_000.0,
        "phiNn_cg_kN": phiNn_cg_N / 1_000.0,
        "interaction_concrete": interaction_concrete,
        "interaction_concrete_ok": interaction_concrete_ok,
        "member_thickness_for_shear_mm": member_thickness_for_shear_mm,
        "thickness_warning": thickness_warning,
        "group_span_perp_mm": group_span_perp_mm,
    }

# ============================================================
# MÓDULO 8 - ACI 17.9 GEOMETRÍA MÍNIMA
# ============================================================

def module8_geometry_minimums_aci_17_9(
    anchors: AnchorLayout,
    pedestal: PedestalGeometry,
    bolt_df: pd.DataFrame,
    anchor_installation: str,
    post_installed_type_for_17_9: str,
    anchor_torqued: bool,
    supplementary_reinforcement_for_splitting: bool,
    nominal_max_agg_mm: float,
    required_cover_mm: float,
    product_specific_geometry_data: bool,
    product_specific_min_edge_mm: float,
    product_specific_min_spacing_mm: float,
    tests_permit_greater_hef: bool,
) -> dict:
    """
    Módulo 8:
    chequeo geométrico mínimo conforme a ACI 318-25 17.9.

    Revisa:
    - espaciamiento mínimo
    - distancia mínima al borde
    - límite de hef para post-installed expansion / screw / undercut

    Notas:
    - Si hay refuerzo suplementario para controlar splitting, 17.9 permite salir de
      estos mínimos generales.
    - Para post-installed, la fuente primaria debe ser información product-specific
      ACI 355.2 / 355.4.
    """

    db = anchors.db_mm
    hef = anchors.hef_mm

    # --------------------------------------------------------
    # Espaciamiento real mínimo entre pernos
    # --------------------------------------------------------
    coords = bolt_df[["x_mm", "y_mm"]].to_numpy()
    min_spacing_real_mm = float("inf")

    for i in range(len(coords)):
        for j in range(i + 1, len(coords)):
            dij = math.dist(coords[i], coords[j])
            min_spacing_real_mm = min(min_spacing_real_mm, dij)

    if min_spacing_real_mm == float("inf"):
        min_spacing_real_mm = 0.0

    # --------------------------------------------------------
    # Distancia real mínima al borde del pedestal
    # --------------------------------------------------------
    x_left_edge = -pedestal.B_ped_mm / 2.0
    x_right_edge = pedestal.B_ped_mm / 2.0
    y_bottom_edge = -pedestal.N_ped_mm / 2.0
    y_top_edge = pedestal.N_ped_mm / 2.0

    edge_distances = []

    for _, row in bolt_df.iterrows():
        x = row["x_mm"]
        y = row["y_mm"]

        edge_distances.extend([
            abs(x - x_left_edge),
            abs(x_right_edge - x),
            abs(y - y_bottom_edge),
            abs(y_top_edge - y),
        ])

    min_edge_real_mm = min(edge_distances) if edge_distances else 0.0

    # --------------------------------------------------------
    # Si existe refuerzo suplementario, el módulo lo reporta,
    # pero igual calcula los mínimos de referencia
    # --------------------------------------------------------
    # Espaciamiento mínimo requerido
    if anchor_installation == "cast_in":
        if anchor_torqued:
            s_min_req_mm = 6.0 * db
            c_min_req_mm = 6.0 * db
            geometry_rule_label = "Cast-in torqued"
        else:
            s_min_req_mm = 4.0 * db
            c_min_req_mm = required_cover_mm
            geometry_rule_label = "Cast-in not torqued"

    elif anchor_installation == "post_installed":
        geometry_rule_label = f"Post-installed: {post_installed_type_for_17_9}"

        if post_installed_type_for_17_9 in ["adhesive", "torque_controlled", "displacement_controlled", "undercut"]:
            s_min_req_mm = 6.0 * db
        elif post_installed_type_for_17_9 == "screw":
            s_min_req_mm = max(0.6 * hef, 6.0 * db)
        else:
            raise ValueError("Tipo de post-installed no reconocido para 17.9.")

        # Distancia al borde
        if product_specific_geometry_data and product_specific_min_edge_mm > 0:
            product_edge_req = product_specific_min_edge_mm
        else:
            if post_installed_type_for_17_9 == "torque_controlled":
                product_edge_req = 8.0 * db
            elif post_installed_type_for_17_9 == "displacement_controlled":
                product_edge_req = 10.0 * db
            elif post_installed_type_for_17_9 in ["screw", "undercut", "adhesive"]:
                product_edge_req = 6.0 * db
            else:
                raise ValueError("Tipo de post-installed no reconocido para borde 17.9.")

        c_min_req_mm = max(
            required_cover_mm,
            2.0 * nominal_max_agg_mm,
            product_edge_req
        )

        if product_specific_geometry_data and product_specific_min_spacing_mm > 0:
            s_min_req_mm = max(s_min_req_mm, product_specific_min_spacing_mm)

    else:
        raise ValueError("anchor_installation debe ser 'cast_in' o 'post_installed'.")

    # --------------------------------------------------------
    # Chequeos
    # --------------------------------------------------------
    spacing_ok = min_spacing_real_mm >= s_min_req_mm
    edge_ok = min_edge_real_mm >= c_min_req_mm

    # --------------------------------------------------------
    # 17.9.4 - límite de hef para ciertos post-installed
    # --------------------------------------------------------
    hef_limit_req_mm = None
    hef_limit_ok = None

    if anchor_installation == "post_installed" and post_installed_type_for_17_9 in [
        "torque_controlled", "displacement_controlled", "screw", "undercut"
    ]:
        hef_limit_req_mm = max((2.0 / 3.0) * pedestal.h_ped_mm, pedestal.h_ped_mm - 100.0)

        if tests_permit_greater_hef:
            hef_limit_ok = True
        else:
            hef_limit_ok = hef <= hef_limit_req_mm

    # --------------------------------------------------------
    # Caso 17.9.3
    # --------------------------------------------------------
    can_use_da_prime = (
        anchor_installation == "cast_in" and
        not anchor_torqued
    ) or (
        anchor_installation == "post_installed" and
        not anchor_torqued
    )

    da_prime_mm = None

    if can_use_da_prime:
        # diámetro equivalente por espaciamiento
        da_prime_spacing = min_spacing_real_mm / 4.0 if anchor_installation == "cast_in" else min_spacing_real_mm / 6.0

        # diámetro equivalente por borde
        if anchor_installation == "cast_in":
            da_prime_edge = float("inf")  # borde controlado por cover, no por múltiplo de da
        else:
            if post_installed_type_for_17_9 == "torque_controlled":
                da_prime_edge = min_edge_real_mm / 8.0
            elif post_installed_type_for_17_9 == "displacement_controlled":
                da_prime_edge = min_edge_real_mm / 10.0
            elif post_installed_type_for_17_9 in ["screw", "undercut", "adhesive"]:
                da_prime_edge = min_edge_real_mm / 6.0
            else:
                da_prime_edge = float("inf")

        da_prime_mm = min(db, da_prime_spacing, da_prime_edge)

    # --------------------------------------------------------
    # Resultado global
    # --------------------------------------------------------
    geometric_ok = spacing_ok and edge_ok

    if hef_limit_ok is not None:
        geometric_ok = geometric_ok and hef_limit_ok

    return {
        "geometry_rule_label": geometry_rule_label,
        "supplementary_reinforcement_for_splitting": supplementary_reinforcement_for_splitting,
        "db_mm": db,
        "hef_mm": hef,
        "min_spacing_real_mm": min_spacing_real_mm,
        "min_edge_real_mm": min_edge_real_mm,
        "s_min_req_mm": s_min_req_mm,
        "c_min_req_mm": c_min_req_mm,
        "spacing_ok": spacing_ok,
        "edge_ok": edge_ok,
        "hef_limit_req_mm": hef_limit_req_mm,
        "hef_limit_ok": hef_limit_ok,
        "can_use_da_prime": can_use_da_prime,
        "da_prime_mm": da_prime_mm,
        "geometric_ok": geometric_ok,
        "required_cover_mm": required_cover_mm,
        "nominal_max_agg_mm": nominal_max_agg_mm,
        "product_specific_geometry_data": product_specific_geometry_data,
        "product_specific_min_edge_mm": product_specific_min_edge_mm,
        "product_specific_min_spacing_mm": product_specific_min_spacing_mm,
        "tests_permit_greater_hef": tests_permit_greater_hef,
    }

# ============================================================
# MÓDULO 9 - TRANSFERENCIA DE CORTANTE EN LA BASE
# VERSIÓN CORREGIDA CON DETECCIÓN AUTOMÁTICA DE SHEAR KEY
# ============================================================

def module9_base_shear_transfer(
    loads: Loads,
    module5_results: dict,
    analysis_mode: str,
    uniaxial_axis: str,
    base_shear_mode: str,
    base_shear_mechanism: str,
    mu: float,
    phi_base_friction: float,
    phiVn_shear_key_external_kN: float,
    allow_combined_mechanisms: bool,
) -> dict:
    """
    Módulo 9:
    transferencia de cortante en la base.

    Modos:
    - auto_detect:
        revisa fricción, anclajes y combinación permitida;
        si no alcanza, concluye que se requiere shear key.
    - manual:
        usa el mecanismo elegido por el usuario.

    Alcance:
    - la shear key entra aquí como necesidad detectada o como
      resistencia ya conocida externamente
    - este módulo NO diseña todavía el acero ni la soldadura del lug
    """

    if analysis_mode != "Uniaxial":
        return None

    # --------------------------------------------------------
    # Cortante demandado
    # --------------------------------------------------------
    if uniaxial_axis == "x":
        Vu_kN = abs(loads.Vux_kN)
    elif uniaxial_axis == "y":
        Vu_kN = abs(loads.Vuy_kN)
    else:
        raise ValueError("El eje uniaxial debe ser 'x' o 'y'.")

    # --------------------------------------------------------
    # Resistencias disponibles
    # --------------------------------------------------------
    Pu_comp_kN = max(loads.Pu_kN, 0.0)
    phiVn_friction_kN = phi_base_friction * mu * Pu_comp_kN
    phiVn_anchor_group_kN = module5_results["phiVsa_kN"] * module5_results["total_bolts"]
    phiVn_shear_key_kN = phiVn_shear_key_external_kN

    mechanism_warning = None
    selected_case = None
    contributing_mechanisms = []
    phiVn_selected_kN = 0.0
    shear_key_required = False
    Vu_remaining_for_key_kN = 0.0

    # --------------------------------------------------------
    # MODO AUTOMÁTICO
    # --------------------------------------------------------
    if base_shear_mode == "auto_detect":

        # 1) Fricción sola
        if Vu_kN <= phiVn_friction_kN:
            selected_case = "friction"
            contributing_mechanisms = ["friction"]
            phiVn_selected_kN = phiVn_friction_kN
            shear_key_required = False

        # 2) Anclajes solos
        elif Vu_kN <= phiVn_anchor_group_kN:
            selected_case = "anchors"
            contributing_mechanisms = ["anchors"]
            phiVn_selected_kN = phiVn_anchor_group_kN
            shear_key_required = False

        # 3) Fricción + anclajes, solo si se permite
        elif allow_combined_mechanisms and Vu_kN <= (phiVn_friction_kN + phiVn_anchor_group_kN):
            selected_case = "friction+anchors"
            contributing_mechanisms = ["friction", "anchors"]
            phiVn_selected_kN = phiVn_friction_kN + phiVn_anchor_group_kN
            shear_key_required = False
            mechanism_warning = (
                "Se está usando combinación fricción + anclajes. "
                "Debes asegurarte de que sea compatible en rigidez y deformación."
            )

        # 4) No alcanza sin shear key
        else:
            selected_case = "shear_key_required"
            shear_key_required = True

            if allow_combined_mechanisms:
                phiVn_selected_kN = phiVn_friction_kN + phiVn_anchor_group_kN
                contributing_mechanisms = ["friction", "anchors"]
            else:
                # tomamos el mejor de los dos mecanismos individuales
                if phiVn_friction_kN >= phiVn_anchor_group_kN:
                    phiVn_selected_kN = phiVn_friction_kN
                    contributing_mechanisms = ["friction"]
                else:
                    phiVn_selected_kN = phiVn_anchor_group_kN
                    contributing_mechanisms = ["anchors"]

            Vu_remaining_for_key_kN = max(Vu_kN - phiVn_selected_kN, 0.0)

            mechanism_warning = (
                "La resistencia disponible sin shear key no alcanza. "
                "Se requiere diseñar una shear key / shear lug para el cortante remanente."
            )

    # --------------------------------------------------------
    # MODO MANUAL
    # --------------------------------------------------------
    elif base_shear_mode == "manual":

        if base_shear_mechanism == "friction":
            phiVn_selected_kN = phiVn_friction_kN
            contributing_mechanisms = ["friction"]
            selected_case = "manual_friction"

        elif base_shear_mechanism == "anchors":
            phiVn_selected_kN = phiVn_anchor_group_kN
            contributing_mechanisms = ["anchors"]
            selected_case = "manual_anchors"

        elif base_shear_mechanism == "shear_key":
            phiVn_selected_kN = phiVn_shear_key_kN
            contributing_mechanisms = ["shear_key"]
            selected_case = "manual_shear_key"

        elif base_shear_mechanism == "combined":
            if not allow_combined_mechanisms:
                phiVn_selected_kN = 0.0
                contributing_mechanisms = []
                selected_case = "manual_combined_not_allowed"
                mechanism_warning = (
                    "Seleccionaste mecanismo combinado, pero no activaste la casilla "
                    "'Permitir combinación fricción + anclajes'."
                )
            else:
                phiVn_selected_kN = (
                    phiVn_friction_kN
                    + phiVn_anchor_group_kN
                    + phiVn_shear_key_kN
                )
                contributing_mechanisms = ["friction", "anchors", "shear_key"]
                selected_case = "manual_combined"
                mechanism_warning = (
                    "La combinación de mecanismos se está usando solo como modelo preliminar. "
                    "Debe asegurarse compatibilidad de rigidez y deformación."
                )
        else:
            raise ValueError("Mecanismo manual no reconocido.")

        shear_key_required = False
        Vu_remaining_for_key_kN = 0.0

    else:
        raise ValueError("base_shear_mode debe ser 'auto_detect' o 'manual'.")

    shear_ok = Vu_kN <= phiVn_selected_kN if phiVn_selected_kN > 0 else False
    utilization = Vu_kN / phiVn_selected_kN if phiVn_selected_kN > 0 else float("inf")

    return {
        "uniaxial_axis": uniaxial_axis,
        "Vu_kN": Vu_kN,
        "base_shear_mode": base_shear_mode,
        "base_shear_mechanism": base_shear_mechanism,
        "selected_case": selected_case,
        "mu": mu,
        "phi_base_friction": phi_base_friction,
        "Pu_comp_kN": Pu_comp_kN,
        "phiVn_friction_kN": phiVn_friction_kN,
        "phiVn_anchor_group_kN": phiVn_anchor_group_kN,
        "phiVn_shear_key_kN": phiVn_shear_key_kN,
        "phiVn_selected_kN": phiVn_selected_kN,
        "contributing_mechanisms": contributing_mechanisms,
        "mechanism_warning": mechanism_warning,
        "shear_key_required": shear_key_required,
        "Vu_remaining_for_key_kN": Vu_remaining_for_key_kN,
        "shear_ok": shear_ok,
        "utilization": utilization,
    }
# ============================================================
# AUXILIAR MÓDULO 10 - LONGITUD EFECTIVA DE SOLDADURA
# ============================================================

def compute_column_weld_length_effective(
    column_plot: dict,
    weld_size_mm: float,
    weld_layout: str,
    manual_length_mm: float,
) -> dict:
    """
    Calcula la longitud efectiva de soldadura columna-placa
    en función de la forma del perfil y del patrón de soldadura.

    Además aplica ajustes básicos de AISC:
    - longitud mínima: 4w
    - reducción por longitud excesiva J2-1
    - si l > 300w, l_eff = 180w
    """

    section_type = column_plot.get("section_type", "W").upper()
    d = column_plot["d_mm"]
    bf = column_plot["bf_mm"]
    tf = column_plot.get("tf_mm", 0.0)

    # --------------------------------------------------------
    # Longitud geométrica "real" elegida
    # --------------------------------------------------------
    if weld_layout == "manual":
        L_geom_mm = manual_length_mm
        layout_used = "manual"

    elif weld_layout == "all_perimeter_rect":
        L_geom_mm = 2.0 * (bf + d)
        layout_used = "all_perimeter_rect"

    elif section_type in ["W", "H", "I"]:
        web_clear_mm = max(d - 2.0 * tf, 0.0)

        if weld_layout == "auto_by_section":
            # por defecto para perfil I: alas + alma accesible
            L_geom_mm = 2.0 * bf + 2.0 * web_clear_mm
            layout_used = "flanges_plus_web_auto"

        elif weld_layout == "flanges_only":
            L_geom_mm = 2.0 * bf
            layout_used = "flanges_only"

        elif weld_layout == "web_only":
            L_geom_mm = 2.0 * web_clear_mm
            layout_used = "web_only"

        elif weld_layout == "flanges_plus_web":
            L_geom_mm = 2.0 * bf + 2.0 * web_clear_mm
            layout_used = "flanges_plus_web"

        else:
            # fallback
            L_geom_mm = 2.0 * bf + 2.0 * web_clear_mm
            layout_used = "fallback_flanges_plus_web"

    else:
        # HSS / BOX / fallback
        if weld_layout == "auto_by_section":
            L_geom_mm = 2.0 * (bf + d)
            layout_used = "all_perimeter_rect_auto"
        elif weld_layout in ["flanges_only", "web_only", "flanges_plus_web"]:
            # no tiene sentido directo en HSS, se redirige
            L_geom_mm = 2.0 * (bf + d)
            layout_used = "all_perimeter_rect_auto_redirect"
        else:
            L_geom_mm = 2.0 * (bf + d)
            layout_used = "fallback_rect"

    if L_geom_mm <= 0:
        raise ValueError("La longitud geométrica de soldadura resultó no positiva.")

    # --------------------------------------------------------
    # AISC J2.2b(c): longitud mínima 4w
    # si no, el tamaño efectivo no excede l/4
    # Para mantener el cálculo en términos de longitud efectiva equivalente:
    # si l < 4w => l_eff = 4*l_actual/w *? no
    # más limpio: reducimos el tamaño efectivo, no la longitud.
    # Pero como el módulo usa Aw = 0.707*w*L_eff,
    # la equivalencia es:
    # w_eff = l/4  => Aw = 0.707*(l/4)*l
    # Para no mezclar tamaños, aquí solo emitimos advertencia.
    # --------------------------------------------------------
    min_length_ok = L_geom_mm >= 4.0 * weld_size_mm

    # --------------------------------------------------------
    # AISC J2-1 para soldaduras cargadas por extremo
    # Se usa aquí como ajuste conservador general para filetes largos.
    # β = 1.2 - 0.002(l/w) <= 1.0
    # Si l > 300w -> l_eff = 180w
    # --------------------------------------------------------
    if weld_size_mm <= 0:
        raise ValueError("El tamaño de soldadura debe ser positivo.")

    l_over_w = L_geom_mm / weld_size_mm

    if l_over_w > 300.0:
        L_eff_mm = 180.0 * weld_size_mm
        beta = L_eff_mm / L_geom_mm
        excessive_length_limit_applied = True
    elif l_over_w > 100.0:
        beta = min(1.0, 1.2 - 0.002 * l_over_w)
        beta = max(beta, 0.0)
        L_eff_mm = beta * L_geom_mm
        excessive_length_limit_applied = True
    else:
        beta = 1.0
        L_eff_mm = L_geom_mm
        excessive_length_limit_applied = False

    return {
        "section_type": section_type,
        "layout_used": layout_used,
        "L_geom_mm": L_geom_mm,
        "L_eff_mm": L_eff_mm,
        "beta_length": beta,
        "l_over_w": l_over_w,
        "min_length_ok": min_length_ok,
        "excessive_length_limit_applied": excessive_length_limit_applied,
    }
# ============================================================
# MÓDULO 10 - SOLDADURA (DEPURADO)
# ============================================================

def module10_weld_design(
    column_plot: dict,
    loads: Loads,
    module3_results: dict,
    module9_results: dict,
    FEXX_MPa: float,
    phi_weld: float,
    column_weld_size_mm: float,
    column_weld_layout: str,
    column_weld_length_manual_mm: float,
    provide_shear_lug_weld: bool,
    shear_lug_weld_size_mm: float,
    shear_lug_weld_length_mm: float,
) -> dict:
    """
    Módulo 10 depurado:
    - soldadura columna-placa según geometría real y patrón elegido
    - soldadura shear lug-placa
    - modelo preliminar de filete cargado concéntricamente

    Base:
    phiRn = phi * 0.60 * FEXX * Aw
    Aw = 0.707 * w * L_eff
    """

    # --------------------------------------------------------
    # Demanda columna-placa
    # --------------------------------------------------------
    if module3_results["tension_active"]:
        Tu_col_base_kN = module3_results["T_total_kN"]
    else:
        Tu_col_base_kN = 0.0

    # --------------------------------------------------------
    # Longitud efectiva columna-placa
    # --------------------------------------------------------
    weld_len_info = compute_column_weld_length_effective(
        column_plot=column_plot,
        weld_size_mm=column_weld_size_mm,
        weld_layout=column_weld_layout,
        manual_length_mm=column_weld_length_manual_mm,
    )

    L_col_geom_mm = weld_len_info["L_geom_mm"]
    L_col_eff_mm = weld_len_info["L_eff_mm"]

    # --------------------------------------------------------
    # Resistencia columna-placa
    # --------------------------------------------------------
    Aw_col_mm2 = 0.707 * column_weld_size_mm * L_col_eff_mm
    phiRn_col_kN = phi_weld * 0.60 * FEXX_MPa * Aw_col_mm2 / 1_000.0

    col_weld_ok = Tu_col_base_kN <= phiRn_col_kN if phiRn_col_kN > 0 else (Tu_col_base_kN == 0.0)
    col_weld_util = Tu_col_base_kN / phiRn_col_kN if phiRn_col_kN > 0 else float("inf")

    # --------------------------------------------------------
    # Demanda lug-placa
    # --------------------------------------------------------
    if module9_results["shear_key_required"]:
        Vu_lug_kN = module9_results["Vu_remaining_for_key_kN"]
    else:
        Vu_lug_kN = 0.0

    # --------------------------------------------------------
    # Resistencia lug-placa
    # --------------------------------------------------------
    if provide_shear_lug_weld:
        if shear_lug_weld_size_mm <= 0 or shear_lug_weld_length_mm <= 0:
            raise ValueError("Para evaluar la soldadura del shear lug, el tamaño y la longitud deben ser positivos.")

        Aw_lug_mm2 = 0.707 * shear_lug_weld_size_mm * shear_lug_weld_length_mm
        phiRn_lug_kN = phi_weld * 0.60 * FEXX_MPa * Aw_lug_mm2 / 1_000.0
        lug_weld_ok = Vu_lug_kN <= phiRn_lug_kN if phiRn_lug_kN > 0 else (Vu_lug_kN == 0.0)
        lug_weld_util = Vu_lug_kN / phiRn_lug_kN if phiRn_lug_kN > 0 else float("inf")
    else:
        Aw_lug_mm2 = 0.0
        phiRn_lug_kN = 0.0
        lug_weld_ok = not module9_results["shear_key_required"]
        lug_weld_util = 0.0 if not module9_results["shear_key_required"] else float("inf")

    return {
        "FEXX_MPa": FEXX_MPa,
        "phi_weld": phi_weld,

        "Tu_col_base_kN": Tu_col_base_kN,
        "column_weld_size_mm": column_weld_size_mm,
        "column_weld_layout": column_weld_layout,
        "column_weld_layout_used": weld_len_info["layout_used"],
        "L_col_geom_mm": L_col_geom_mm,
        "L_col_weld_eff_mm": L_col_eff_mm,
        "beta_length": weld_len_info["beta_length"],
        "l_over_w": weld_len_info["l_over_w"],
        "min_length_ok": weld_len_info["min_length_ok"],
        "excessive_length_limit_applied": weld_len_info["excessive_length_limit_applied"],
        "Aw_col_mm2": Aw_col_mm2,
        "phiRn_col_kN": phiRn_col_kN,
        "col_weld_ok": col_weld_ok,
        "col_weld_util": col_weld_util,

        "shear_key_required": module9_results["shear_key_required"],
        "Vu_lug_kN": Vu_lug_kN,
        "provide_shear_lug_weld": provide_shear_lug_weld,
        "shear_lug_weld_size_mm": shear_lug_weld_size_mm,
        "shear_lug_weld_length_mm": shear_lug_weld_length_mm,
        "Aw_lug_mm2": Aw_lug_mm2,
        "phiRn_lug_kN": phiRn_lug_kN,
        "lug_weld_ok": lug_weld_ok,
        "lug_weld_util": lug_weld_util,
    }

# ============================================================
# MÓDULO 11 - BIAxIAL PRELIMINAR
# ============================================================

def module11_biaxial_preliminary(
    loads: Loads,
    base_plate: BasePlateGeometry,
    bolt_df: pd.DataFrame,
) -> dict:
    """
    Módulo 11:
    análisis preliminar biaxial elástico de presión sobre placa base
    y tracción preliminar en pernos.

    Modelo:
        q(x,y) = P/A + (My/Iy)*x + (Mx/Ix)*y

    Salidas:
    - presión en esquinas
    - qmax, qmin
    - clasificación preliminar
    - tracción preliminar en pernos si qmin < 0
    """

    B = base_plate.B_bp_mm
    N = base_plate.N_bp_mm

    A = B * N
    Ix = B * N**3 / 12.0
    Iy = N * B**3 / 12.0

    Pu_N = loads.Pu_kN * 1_000.0
    Mux_Nmm = loads.Mux_kNm * 1_000_000.0
    Muy_Nmm = loads.Muy_kNm * 1_000_000.0

    # --------------------------------------------------------
    # Presión elástica biaxial
    # --------------------------------------------------------
    def q_biaxial(x_mm: float, y_mm: float) -> float:
        return (Pu_N / A) + (Muy_Nmm / Iy) * x_mm + (Mux_Nmm / Ix) * y_mm

    # Esquinas
    corners = {
        "(+B/2,+N/2)": ( B/2.0,  N/2.0),
        "(+B/2,-N/2)": ( B/2.0, -N/2.0),
        "(-B/2,+N/2)": (-B/2.0,  N/2.0),
        "(-B/2,-N/2)": (-B/2.0, -N/2.0),
    }

    corner_pressures = {}
    for label, (x, y) in corners.items():
        corner_pressures[label] = q_biaxial(x, y)

    q_values = list(corner_pressures.values())
    q_max = max(q_values)
    q_min = min(q_values)

    full_compression = q_min >= 0.0
    possible_uplift = q_min < 0.0

    # --------------------------------------------------------
    # Excentricidades geométricas
    # --------------------------------------------------------
    e_x_mm = Muy_Nmm / Pu_N
    e_y_mm = Mux_Nmm / Pu_N

    # --------------------------------------------------------
    # Tracción preliminar en pernos
    # --------------------------------------------------------
    bolt_tension_df = bolt_df.copy()
    bolt_tension_df["demand_index"] = 0.0
    bolt_tension_df["T_prelim_kN"] = 0.0

    T_total_prelim_kN = 0.0
    critical_bolt = None
    Tmax_prelim_kN = 0.0

    if possible_uplift:
        # índice de demanda elástica preliminar
        # combinación del signo real de Mx y My
        bolt_tension_df["demand_index"] = (
            (Muy_Nmm / Iy) * bolt_tension_df["x_mm"] +
            (Mux_Nmm / Ix) * bolt_tension_df["y_mm"]
        )

        # Solo pernos del lado que más tracciona
        demand_min = bolt_tension_df["demand_index"].min()

        # Shift para que el mínimo sea 0
        bolt_tension_df["demand_pos"] = bolt_tension_df["demand_index"] - demand_min

        total_index = bolt_tension_df["demand_pos"].sum()

        # Tracción total preliminar:
        # se toma como exceso de la presión negativa convertida a fuerza sobre área
        # aproximación global:
        uplift_ratio = abs(q_min) / max(q_max, 1e-9)
        T_total_prelim_kN = uplift_ratio * loads.Pu_kN

        if total_index > 0:
            bolt_tension_df["T_prelim_kN"] = (
                T_total_prelim_kN * bolt_tension_df["demand_pos"] / total_index
            )
        else:
            bolt_tension_df["T_prelim_kN"] = 0.0

        Tmax_prelim_kN = bolt_tension_df["T_prelim_kN"].max()
        if Tmax_prelim_kN > 0:
            critical_bolt_row = bolt_tension_df.loc[bolt_tension_df["T_prelim_kN"].idxmax()]
            critical_bolt = int(critical_bolt_row["Perno"])

    return {
        "A_mm2": A,
        "Ix_mm4": Ix,
        "Iy_mm4": Iy,
        "Pu_kN": loads.Pu_kN,
        "Mux_kNm": loads.Mux_kNm,
        "Muy_kNm": loads.Muy_kNm,
        "e_x_mm": e_x_mm,
        "e_y_mm": e_y_mm,
        "corner_pressures_MPa": corner_pressures,
        "q_max_MPa": q_max,
        "q_min_MPa": q_min,
        "full_compression": full_compression,
        "possible_uplift": possible_uplift,
        "T_total_prelim_kN": T_total_prelim_kN,
        "Tmax_prelim_kN": Tmax_prelim_kN,
        "critical_bolt": critical_bolt,
        "bolt_tension_df": bolt_tension_df,
    }
# ============================================================
# MÓDULO 12 - REFINAMIENTO BIAxIAL POR MALLA
# ============================================================

def module12_biaxial_grid_refinement(
    loads: Loads,
    base_plate: BasePlateGeometry,
    bolt_df: pd.DataFrame,
    nx: int,
    ny: int,
    use_bolt_distribution: bool = True,
) -> dict:
    """
    Módulo 12:
    refinamiento biaxial mediante malla sobre la placa.

    Modelo:
    q(x,y) = P/A + (My/Iy)*x + (Mx/Ix)*y

    Se separa:
    - compresión: q_pos = max(q, 0)
    - levantamiento elástico: q_neg = max(-q, 0)

    Este módulo NO es todavía el equilibrio no lineal final,
    pero mejora notablemente la evaluación preliminar biaxial.
    """

    B = base_plate.B_bp_mm
    N = base_plate.N_bp_mm

    A = B * N
    Ix = B * N**3 / 12.0
    Iy = N * B**3 / 12.0

    Pu_N = loads.Pu_kN * 1_000.0
    Mux_Nmm = loads.Mux_kNm * 1_000_000.0
    Muy_Nmm = loads.Muy_kNm * 1_000_000.0

    # --------------------------------------------------------
    # Malla
    # --------------------------------------------------------
    x_vals = np.linspace(-B/2.0, B/2.0, int(nx))
    y_vals = np.linspace(-N/2.0, N/2.0, int(ny))

    dx = B / (int(nx) - 1)
    dy = N / (int(ny) - 1)
    dA = dx * dy

    X, Y = np.meshgrid(x_vals, y_vals)

    # --------------------------------------------------------
    # Presión elástica biaxial
    # --------------------------------------------------------
    Q = (Pu_N / A) + (Muy_Nmm / Iy) * X + (Mux_Nmm / Ix) * Y

    Qpos = np.maximum(Q, 0.0)
    Qneg = np.maximum(-Q, 0.0)

    q_max_MPa = float(np.max(Q))
    q_min_MPa = float(np.min(Q))

    full_compression = q_min_MPa >= 0.0
    possible_uplift = q_min_MPa < 0.0

    # --------------------------------------------------------
    # Integración de compresión y levantamiento elástico
    # --------------------------------------------------------
    C_N = float(np.sum(Qpos) * dA)
    T_equiv_N = float(np.sum(Qneg) * dA)

    if C_N > 0:
        xC_mm = float(np.sum(Qpos * X) * dA / C_N)
        yC_mm = float(np.sum(Qpos * Y) * dA / C_N)
    else:
        xC_mm = 0.0
        yC_mm = 0.0

    # --------------------------------------------------------
    # Tracción preliminar en pernos
    # --------------------------------------------------------
    bolt_tension_df = bolt_df.copy()
    bolt_tension_df["uplift_index"] = 0.0
    bolt_tension_df["T_refined_kN"] = 0.0

    Tmax_refined_kN = 0.0
    critical_bolt = None

    if possible_uplift and use_bolt_distribution:
        # Índice de levantamiento en pernos usando la misma ley elástica
        q_bolts = (
            (Pu_N / A)
            + (Muy_Nmm / Iy) * bolt_tension_df["x_mm"].to_numpy()
            + (Mux_Nmm / Ix) * bolt_tension_df["y_mm"].to_numpy()
        )

        uplift_index = np.maximum(-q_bolts, 0.0)

        # Si por discretización sale todo cero pero sí hay uplift, se usa fallback
        if np.sum(uplift_index) <= 0:
            fallback = (
                (Muy_Nmm / Iy) * bolt_tension_df["x_mm"].to_numpy()
                + (Mux_Nmm / Ix) * bolt_tension_df["y_mm"].to_numpy()
            )
            fallback = fallback - np.min(fallback)
            uplift_index = np.maximum(fallback, 0.0)

        bolt_tension_df["uplift_index"] = uplift_index

        total_index = float(np.sum(uplift_index))

        if total_index > 0 and T_equiv_N > 0:
            bolt_tension_df["T_refined_kN"] = (
                T_equiv_N * uplift_index / total_index / 1_000.0
            )
        else:
            bolt_tension_df["T_refined_kN"] = 0.0

        Tmax_refined_kN = float(bolt_tension_df["T_refined_kN"].max())

        if Tmax_refined_kN > 0:
            critical_bolt_row = bolt_tension_df.loc[bolt_tension_df["T_refined_kN"].idxmax()]
            critical_bolt = int(critical_bolt_row["Perno"])

    # --------------------------------------------------------
    # Momentos internos aproximados
    # --------------------------------------------------------
    T_i_N = bolt_tension_df["T_refined_kN"].to_numpy() * 1_000.0
    x_i = bolt_tension_df["x_mm"].to_numpy()
    y_i = bolt_tension_df["y_mm"].to_numpy()

    # Convención:
    # Mx por brazo en y
    # My por brazo en x
    Mx_internal_Nmm = C_N * yC_mm - np.sum(T_i_N * y_i)
    My_internal_Nmm = C_N * xC_mm - np.sum(T_i_N * x_i)

    Mx_residual_kNm = loads.Mux_kNm - Mx_internal_Nmm / 1_000_000.0
    My_residual_kNm = loads.Muy_kNm - My_internal_Nmm / 1_000_000.0

    return {
        "A_mm2": A,
        "Ix_mm4": Ix,
        "Iy_mm4": Iy,
        "nx": int(nx),
        "ny": int(ny),
        "dx_mm": dx,
        "dy_mm": dy,
        "dA_mm2": dA,
        "q_max_MPa": q_max_MPa,
        "q_min_MPa": q_min_MPa,
        "full_compression": full_compression,
        "possible_uplift": possible_uplift,
        "C_kN": C_N / 1_000.0,
        "T_equiv_kN": T_equiv_N / 1_000.0,
        "xC_mm": xC_mm,
        "yC_mm": yC_mm,
        "Mx_internal_kNm": Mx_internal_Nmm / 1_000_000.0,
        "My_internal_kNm": My_internal_Nmm / 1_000_000.0,
        "Mx_residual_kNm": Mx_residual_kNm,
        "My_residual_kNm": My_residual_kNm,
        "critical_bolt": critical_bolt,
        "Tmax_refined_kN": Tmax_refined_kN,
        "bolt_tension_df": bolt_tension_df,
        "X": X,
        "Y": Y,
        "Q": Q,
        "Qpos": Qpos,
        "Qneg": Qneg,
    }
# ============================================================
# MÓDULO 13 - CIERRE DEL DISEÑO Y REPORTE GLOBAL
# ============================================================

def module13_design_summary(
    analysis_mode: str,
    module2_results=None,
    module4_results=None,
    module5_results=None,
    module6_results=None,
    module7_results=None,
    module8_results=None,
    module9_results=None,
    module10_results=None,
    module11_results=None,
    module12_results=None,
) -> dict:
    """
    Módulo 13:
    consolidación global de resultados.

    Devuelve:
    - tabla de chequeos
    - estado global
    - lista de módulos críticos
    """

    checks = []

    def add_check(module_name, check_name, status, utilization=None, note=""):
        checks.append({
            "Módulo": module_name,
            "Chequeo": check_name,
            "Estado": status,
            "Utilización": utilization,
            "Nota": note,
        })

    # --------------------------------------------------------
    # RAMA UNIAXIAL
    # --------------------------------------------------------
    if analysis_mode == "Uniaxial":

        if module2_results is not None:
            add_check(
                "Módulo 2",
                "Bearing bajo placa",
                "Cumple" if module2_results["bearing_ok"] else "No cumple",
                utilization=(module2_results["q_max_MPa"] / module2_results["q_allow_phi_MPa"]
                             if module2_results["q_allow_phi_MPa"] > 0 else None),
                note=f"Caso: {module2_results['case']}"
            )

        if module4_results is not None:
            add_check(
                "Módulo 4",
                "Espesor de placa",
                "Cumple" if module4_results["thickness_ok"] else "No cumple",
                utilization=module4_results["utilization"],
                note="tp >= t_req"
            )

        if module5_results is not None:
            add_check(
                "Módulo 5",
                "Acero del anclaje - tensión",
                "Cumple" if module5_results["tension_ok"] else "No cumple",
                utilization=module5_results["tension_ratio"],
            )
            add_check(
                "Módulo 5",
                "Acero del anclaje - cortante",
                "Cumple" if module5_results["shear_ok"] else "No cumple",
                utilization=module5_results["shear_ratio"],
            )
            add_check(
                "Módulo 5",
                "Acero del anclaje - interacción",
                "Cumple" if module5_results["interaction_ok"] else "No cumple",
                utilization=module5_results["interaction_value"],
            )

        if module6_results is not None:
            add_check(
                "Módulo 6",
                "Concreto en tensión",
                "Cumple" if module6_results["concrete_tension_ok"] else "No cumple",
                utilization=(module6_results["Nua_group_kN"] / module6_results["phiNn_cg_kN"]
                             if module6_results["phiNn_cg_kN"] > 0 else None),
                note="ACI 17 tensión"
            )

        if module7_results is not None:
            add_check(
                "Módulo 7",
                "Concreto en cortante",
                "Cumple" if module7_results["shear_concrete_ok"] else "No cumple",
                utilization=(module7_results["Vua_group_kN"] / module7_results["phiVn_cg_kN"]
                             if module7_results["phiVn_cg_kN"] > 0 else None),
                note="ACI 17 cortante"
            )
            add_check(
                "Módulo 7",
                "Interacción concreta N-V",
                "Cumple" if module7_results["interaction_concrete_ok"] else "No cumple",
                utilization=module7_results["interaction_concrete"],
            )

        if module8_results is not None:
            add_check(
                "Módulo 8",
                "Geometría mínima ACI 17.9",
                "Cumple" if module8_results["geometric_ok"] else "No cumple",
                utilization=None,
                note=module8_results["geometry_rule_label"]
            )

        if module9_results is not None:
            add_check(
                "Módulo 9",
                "Mecanismo de cortante en la base",
                "Cumple" if module9_results["shear_ok"] else "No cumple",
                utilization=module9_results["utilization"],
                note=f"Mecanismo: {module9_results['selected_case']}"
            )

            if module9_results["shear_key_required"]:
                add_check(
                    "Módulo 9",
                    "Necesidad de shear key",
                    "Revisión requerida",
                    utilization=None,
                    note=f"Vu remanente = {module9_results['Vu_remaining_for_key_kN']:.3f} kN"
                )

        if module10_results is not None:
            add_check(
                "Módulo 10",
                "Soldadura columna-placa",
                "Cumple" if module10_results["col_weld_ok"] else "No cumple",
                utilization=module10_results["col_weld_util"],
            )

            if module10_results["shear_key_required"]:
                if module10_results["provide_shear_lug_weld"]:
                    add_check(
                        "Módulo 10",
                        "Soldadura shear lug-placa",
                        "Cumple" if module10_results["lug_weld_ok"] else "No cumple",
                        utilization=module10_results["lug_weld_util"],
                    )
                else:
                    add_check(
                        "Módulo 10",
                        "Soldadura shear lug-placa",
                        "Revisión requerida",
                        utilization=None,
                        note="Se requiere shear key pero no se evaluó su soldadura"
                    )

    # --------------------------------------------------------
    # RAMA BIAXIAL
    # --------------------------------------------------------
    elif analysis_mode == "Biaxial":

        if module11_results is not None:
            add_check(
                "Módulo 11",
                "Biaxial preliminar",
                "Revisión requerida" if module11_results["possible_uplift"] else "Preliminarmente aceptable",
                utilization=None,
                note="Modelo elástico por esquinas"
            )

        if module12_results is not None:
            residual_norm = max(
                abs(module12_results["Mx_residual_kNm"]),
                abs(module12_results["My_residual_kNm"]),
            )

            status = "Revisión requerida"
            note = "Refinamiento por malla; aún no es equilibrio no lineal final"

            if not module12_results["possible_uplift"]:
                status = "Preliminarmente aceptable"

            add_check(
                "Módulo 12",
                "Biaxial refinado por malla",
                status,
                utilization=None,
                note=note
            )

            add_check(
                "Módulo 12",
                "Residual de equilibrio biaxial",
                "Revisión requerida" if residual_norm > 1e-3 else "Aceptable",
                utilization=residual_norm,
                note="máx(|Mx_res|, |My_res|) [kN·m]"
            )

    # --------------------------------------------------------
    # ESTADO GLOBAL
    # --------------------------------------------------------
    status_priority = {
        "Cumple": 0,
        "Preliminarmente aceptable": 1,
        "Revisión requerida": 2,
        "No cumple": 3,
    }

    if not checks:
        global_status = "Sin resultados"
        critical_checks = []
    else:
        worst = max(checks, key=lambda row: status_priority.get(row["Estado"], 99))
        global_status = worst["Estado"]

        critical_checks = [
            row for row in checks
            if row["Estado"] in ["No cumple", "Revisión requerida"]
        ]

    summary_df = pd.DataFrame(checks)

    return {
        "summary_df": summary_df,
        "global_status": global_status,
        "critical_checks": critical_checks,
    }
# ============================================================
# MÓDULO 14 - GRÁFICAS DE PRESIÓN Y TRACCIÓN
# ============================================================

def _uniaxial_contact_distribution(module2_results, base_plate: BasePlateGeometry, axis: str, npts: int = 201):
    """
    Devuelve coordenadas y presión de contacto uniaxial para graficar.
    """

    if axis == "x":
        L = base_plate.N_bp_mm
    elif axis == "y":
        L = base_plate.B_bp_mm
    else:
        raise ValueError("axis debe ser 'x' o 'y'.")

    s = np.linspace(-L/2.0, L/2.0, int(npts))
    q = np.zeros_like(s)

    e = module2_results["e_mm"]
    qmax = module2_results["q_max_MPa"]

    if module2_results["case"] == "full_compression":
        qavg = module2_results["q_avg_MPa"]
        factor = 6.0 * e / L
        q = qavg * (1.0 + 2.0 * factor * s / (L/2.0))
        q = np.maximum(q, 0.0)

    else:
        a = module2_results["a_comp_mm"]

        if e >= 0:
            s0 = L/2.0 - a
            s1 = L/2.0
            mask = (s >= s0) & (s <= s1)
            q[mask] = qmax * (s[mask] - s0) / a
        else:
            s0 = -L/2.0
            s1 = -L/2.0 + a
            mask = (s >= s0) & (s <= s1)
            q[mask] = qmax * (s1 - s[mask]) / a

    return s, q


def plot_uniaxial_pressure(module2_results, base_plate: BasePlateGeometry, axis: str, npts: int = 201):
    """
    Gráfica uniaxial de presión de contacto.
    """
    s, q = _uniaxial_contact_distribution(module2_results, base_plate, axis, npts=npts)

    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(s, q, linewidth=2)
    ax.fill_between(s, q, 0, alpha=0.25)

    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_xlabel(f"Coordenada en dirección resistente {'N' if axis == 'x' else 'B'} [mm]")
    ax.set_ylabel("Presión de contacto [MPa]")
    ax.set_title("Distribución uniaxial de presión de contacto")
    ax.grid(True, linestyle="--", alpha=0.35)

    return fig


def plot_uniaxial_anchor_tension(base_plate: BasePlateGeometry, bolt_df: pd.DataFrame, module3_results: dict, axis: str):
    """
    Mapa en planta de tracción en pernos para el caso uniaxial.
    """
    B = base_plate.B_bp_mm
    N = base_plate.N_bp_mm

    df = bolt_df.copy()
    df["T_plot_kN"] = 0.0

    if module3_results["tension_active"]:
        critical_bolts = module3_results["critical_bolts"]
        Tbolt = module3_results["T_per_bolt_kN"]
        df.loc[df["Perno"].isin(critical_bolts), "T_plot_kN"] = Tbolt

    fig, ax = plt.subplots(figsize=(7, 7))

    # placa
    ax.add_patch(Rectangle(
        (-B/2, -N/2), B, N,
        fill=False, linewidth=2.0, edgecolor="black"
    ))

    vmax = max(df["T_plot_kN"].max(), 1e-9)

    for _, row in df.iterrows():
        sc = ax.scatter(
            row["x_mm"], row["y_mm"],
            s=180,
            c=row["T_plot_kN"],
            vmin=0.0,
            vmax=vmax,
            cmap="Reds",
            edgecolors="blue"
        )
        ax.text(row["x_mm"] + 6, row["y_mm"] + 6, str(int(row["Perno"])), fontsize=9)

    cbar = fig.colorbar(sc, ax=ax)
    cbar.set_label("Tracción por perno [kN]")

    ax.set_aspect("equal", adjustable="box")
    ax.set_xlabel("x [mm]")
    ax.set_ylabel("y [mm]")
    ax.set_title("Tracción uniaxial en pernos")
    ax.grid(True, linestyle="--", alpha=0.35)

    return fig


def plot_biaxial_pressure_field(module12_results: dict, title: str = "Presión biaxial", use_positive_only: bool = False):
    """
    Mapa de presión biaxial usando la malla del Módulo 12.
    """
    X = module12_results["X"]
    Y = module12_results["Y"]

    if use_positive_only:
        Z = module12_results["Qpos"]
    else:
        Z = module12_results["Q"]

    fig, ax = plt.subplots(figsize=(7, 6))
    m = ax.contourf(X, Y, Z, levels=30, cmap="coolwarm")
    fig.colorbar(m, ax=ax, label="Presión [MPa]")

    ax.set_aspect("equal", adjustable="box")
    ax.set_xlabel("x [mm]")
    ax.set_ylabel("y [mm]")
    ax.set_title(title)
    ax.grid(True, linestyle="--", alpha=0.20)

    return fig


def plot_biaxial_anchor_tension(base_plate: BasePlateGeometry, module12_results: dict):
    """
    Mapa en planta de tracción refinada en pernos para el caso biaxial.
    """
    B = base_plate.B_bp_mm
    N = base_plate.N_bp_mm

    df = module12_results["bolt_tension_df"].copy()

    fig, ax = plt.subplots(figsize=(7, 7))

    # placa
    ax.add_patch(Rectangle(
        (-B/2, -N/2), B, N,
        fill=False, linewidth=2.0, edgecolor="black"
    ))

    vmax = max(df["T_refined_kN"].max(), 1e-9)

    for _, row in df.iterrows():
        sc = ax.scatter(
            row["x_mm"], row["y_mm"],
            s=180,
            c=row["T_refined_kN"],
            vmin=0.0,
            vmax=vmax,
            cmap="Reds",
            edgecolors="blue"
        )
        ax.text(row["x_mm"] + 6, row["y_mm"] + 6, str(int(row["Perno"])), fontsize=9)

    cbar = fig.colorbar(sc, ax=ax)
    cbar.set_label("Tracción refinada por perno [kN]")

    # centroide de compresión
    ax.plot(module12_results["xC_mm"], module12_results["yC_mm"], "k+", markersize=12, markeredgewidth=2)

    ax.set_aspect("equal", adjustable="box")
    ax.set_xlabel("x [mm]")
    ax.set_ylabel("y [mm]")
    ax.set_title("Tracción refinada biaxial en pernos")
    ax.grid(True, linestyle="--", alpha=0.35)

    return fig

# ============================================================
# MÓDULO 15 - AUXILIARES PARA MEMORIA WORD
# ============================================================

def docx_add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)


def docx_add_subtitle(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)


def docx_add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def docx_add_equation_line(doc: Document, eq_text: str):
    """
    Primera versión estable:
    agrega ecuación como línea centrada y formateada.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(eq_text)
    r.italic = True
    r.font.size = Pt(11)


def docx_add_normal_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def docx_add_small_note(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)


def docx_add_key_value_table(doc: Document, rows_data: list, title: str = None):
    """
    rows_data = [(parametro, valor), ...]
    """
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Parámetro"
    hdr[1].text = "Valor"

    for k, v in rows_data:
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)

    doc.add_paragraph("")


def docx_add_image(doc: Document, image_path: str, caption: str = None, width_inches: float = 5.8):
    """
    Inserta imagen centrada con caption.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))

    if caption:
        pcap = doc.add_paragraph()
        pcap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pcap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)

def get_graph_interpretation(graph_key: str, analysis_mode: str, module2_results=None, module11_results=None, module12_results=None):
    """
    Devuelve un texto corto interpretando la gráfica.
    """

    if graph_key == "uniaxial_pressure":
        if module2_results is not None:
            if module2_results["case"] == "full_compression":
                return (
                    "La gráfica muestra la distribución uniaxial de presión de contacto bajo la placa. "
                    "Se observa que toda la longitud analizada permanece en compresión, por lo que no se "
                    "detecta levantamiento en este caso uniaxial."
                )
            else:
                return (
                    "La gráfica muestra una distribución uniaxial con compresión parcial. "
                    "Solo una parte de la placa permanece en contacto efectivo con el concreto, "
                    "lo que indica levantamiento en la zona opuesta."
                )

    elif graph_key == "uniaxial_anchor_tension":
        return (
            "La gráfica muestra la distribución de tracción en los pernos para el análisis uniaxial. "
            "Los pernos con mayor intensidad corresponden a la fila extrema del lado traccionado."
        )

    elif graph_key == "biaxial_elastic_pressure":
        if module11_results is not None:
            if module11_results["possible_uplift"]:
                return (
                    "El mapa representa la presión biaxial elástica q(x,y). "
                    "Las zonas con presión negativa indican tendencia al levantamiento de la placa."
                )
            else:
                return (
                    "El mapa representa la presión biaxial elástica q(x,y). "
                    "Toda la placa permanece en compresión dentro de esta evaluación preliminar."
                )

    elif graph_key == "biaxial_contact_pressure":
        return (
            "El mapa muestra únicamente la presión de contacto positiva q⁺(x,y), es decir, "
            "la parte de la placa que realmente comprime el concreto. "
            "Las regiones en blanco corresponden a zonas sin contacto efectivo."
        )

    elif graph_key == "biaxial_anchor_tension":
        if module12_results is not None:
            if module12_results["possible_uplift"]:
                return (
                    "La gráfica muestra la distribución refinada de tracción en el grupo de pernos "
                    "para el caso biaxial. Los pernos con mayor intensidad son los más demandados "
                    "por el levantamiento estimado de la placa."
                )
            else:
                return (
                    "La gráfica muestra la distribución refinada de tracción en pernos para el caso biaxial. "
                    "En este caso no se observa levantamiento significativo, por lo que la tracción en pernos es reducida o nula."
                )

    return "La figura presenta el resultado gráfico asociado al módulo de análisis correspondiente."

def save_figure_to_temp(fig, filename_stub: str):
    tmp_dir = tempfile.gettempdir()
    path = os.path.join(tmp_dir, f"{filename_stub}.png")
    fig.savefig(path, dpi=200, bbox_inches="tight")
    return path

# ============================================================
# MÓDULO 15 - GENERACIÓN DE MEMORIA DE CÁLCULO EN WORD
# ============================================================

def module15_generate_word_report(
    analysis_mode: str,
    loads: Loads,
    materials: Materials,
    column_plot: dict,
    base_plate: BasePlateGeometry,
    anchors: AnchorLayout,
    pedestal: PedestalGeometry,
    module13_results: dict,
    module14_figs: dict,
    module2_results=None,
    module4_results=None,
    module5_results=None,
    module6_results=None,
    module7_results=None,
    module8_results=None,
    module9_results=None,
    module10_results=None,
    module11_results=None,
    module12_results=None,
):
    """
    Genera la memoria de cálculo en formato .docx.
    """

    doc = Document()

    # Márgenes
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    # --------------------------------------------------------
    # Portada
    # --------------------------------------------------------
    docx_add_title(doc, "MEMORIA DE CÁLCULO")
    docx_add_title(doc, "DISEÑO DE PLACA BASE DE ACERO")
    docx_add_subtitle(doc, f"Fecha de generación: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    docx_add_subtitle(doc, "Base normativa: AISC 360-22, ACI 318-25")
    doc.add_paragraph("")

    docx_add_normal_paragraph(
        doc,
        "El presente documento resume el proceso de verificación de una conexión de placa base de acero "
        "sobre pedestal de hormigón, incluyendo revisión de presión bajo placa, espesor de placa, "
        "anclajes, concreto, cortante en la base y soldadura, según el modo de análisis seleccionado."
    )

    doc.add_page_break()

    # --------------------------------------------------------
    # 1. Datos de entrada
    # --------------------------------------------------------
    docx_add_heading(doc, "1. Datos de entrada", level=1)

    docx_add_key_value_table(doc, [
        ("Pu [kN]", f"{loads.Pu_kN:.3f}"),
        ("Mux [kN·m]", f"{loads.Mux_kNm:.3f}"),
        ("Muy [kN·m]", f"{loads.Muy_kNm:.3f}"),
        ("Vux [kN]", f"{loads.Vux_kN:.3f}"),
        ("Vuy [kN]", f"{loads.Vuy_kN:.3f}"),
    ], title="Cargas")

    docx_add_key_value_table(doc, [
        ("Fy placa [MPa]", f"{materials.Fy_plate_MPa:.3f}"),
        ("Fu placa [MPa]", f"{materials.Fu_plate_MPa:.3f}"),
        ("f'c [MPa]", f"{materials.fc_MPa:.3f}"),
        ("Fy anclaje [MPa]", f"{materials.Fy_anchor_MPa:.3f}"),
        ("Fu anclaje [MPa]", f"{materials.Fu_anchor_MPa:.3f}"),
    ], title="Materiales")

    docx_add_key_value_table(doc, [
        ("Tipo de perfil", str(column_plot.get("section_type", "-"))),
        ("d [mm]", f"{column_plot.get('d_mm', 0.0):.3f}"),
        ("bf [mm]", f"{column_plot.get('bf_mm', 0.0):.3f}"),
        ("tf [mm]", f"{column_plot.get('tf_mm', 0.0):.3f}"),
        ("tw [mm]", f"{column_plot.get('tw_mm', 0.0):.3f}"),
    ], title="Perfil de columna")

    docx_add_key_value_table(doc, [
        ("B placa [mm]", f"{base_plate.B_bp_mm:.3f}"),
        ("N placa [mm]", f"{base_plate.N_bp_mm:.3f}"),
        ("tp [mm]", f"{base_plate.tp_mm:.3f}"),
    ], title="Placa base")

    docx_add_key_value_table(doc, [
        ("nbx", f"{anchors.nbx}"),
        ("nby", f"{anchors.nby}"),
        ("edge_x [mm]", f"{anchors.edge_x_mm:.3f}"),
        ("edge_y [mm]", f"{anchors.edge_y_mm:.3f}"),
        ("db [mm]", f"{anchors.db_mm:.3f}"),
        ("Ab [mm²]", f"{anchors.Ab_mm2:.3f}"),
        ("hef [mm]", f"{anchors.hef_mm:.3f}"),
    ], title="Anclajes")

    docx_add_key_value_table(doc, [
        ("B pedestal [mm]", f"{pedestal.B_ped_mm:.3f}"),
        ("N pedestal [mm]", f"{pedestal.N_ped_mm:.3f}"),
        ("h pedestal [mm]", f"{pedestal.h_ped_mm:.3f}"),
    ], title="Pedestal")

    # --------------------------------------------------------
    # 2. Desarrollo del cálculo
    # --------------------------------------------------------
    docx_add_heading(doc, "2. Desarrollo del cálculo", level=1)

    if analysis_mode == "Uniaxial":
        # Módulo 2
        if module2_results is not None:
            docx_add_heading(doc, "2.1 Presión de contacto bajo placa", level=2)
            docx_add_normal_paragraph(
                doc,
                "La presión de contacto bajo placa se evaluó bajo hipótesis uniaxial. "
                "Cuando la resultante permanece dentro del núcleo central, se adopta distribución lineal de presiones. "
                "Si la excentricidad excede el núcleo, se considera compresión parcial."
            )
            docx_add_equation_line(doc, "q_max = (P/A) · (1 + 6e/L)")
            docx_add_equation_line(doc, "q_min = (P/A) · (1 - 6e/L)")
            docx_add_small_note(
                doc,
                "Referencia: AISC 360-22, Sección J8 (Column Bases and Bearing on Concrete)."
            )
            docx_add_key_value_table(doc, [
                ("Caso", "Compresión total" if module2_results["case"] == "full_compression" else "Compresión parcial"),
                ("q_max [MPa]", f"{module2_results['q_max_MPa']:.5f}"),
                ("q_min [MPa]", f"{module2_results['q_min_MPa']:.5f}"),
                ("a comprimida [mm]", f"{module2_results['a_comp_mm']:.3f}"),
                ("φ·0.85·f'c [MPa]", f"{module2_results['q_allow_phi_MPa']:.5f}"),
                ("Bearing", "Cumple" if module2_results["bearing_ok"] else "No cumple"),
            ])

        # Módulo 4
        if module4_results is not None:
            docx_add_heading(doc, "2.2 Espesor de placa", level=2)
            docx_add_normal_paragraph(
                doc,
                "El espesor mínimo requerido se estimó modelando la placa como una franja en voladizo "
                "cargada por la presión de contacto gobernante."
            )
            docx_add_equation_line(doc, "t_req = m · √(2q_u / (φFy))")
            docx_add_small_note(
                doc,
                "Referencia: formulación práctica para placa base en flexión local dentro del marco de AISC 360-22 J8."
            )
            docx_add_key_value_table(doc, [
                ("mx [mm]", f"{module4_results['mx_mm']:.3f}"),
                ("my [mm]", f"{module4_results['my_mm']:.3f}"),
                ("m crítico [mm]", f"{module4_results['mcrit_mm']:.3f}"),
                ("q_u [MPa]", f"{module4_results['q_u_MPa']:.5f}"),
                ("tp requerido [mm]", f"{module4_results['t_req_mm']:.3f}"),
                ("tp adoptado [mm]", f"{module4_results['tp_input_mm']:.3f}"),
                ("Chequeo", "Cumple" if module4_results["thickness_ok"] else "No cumple"),
            ])

        # Módulo 5
        if module5_results is not None:
            docx_add_heading(doc, "2.3 Acero del anclaje", level=2)
            docx_add_normal_paragraph(
                doc,
                "Se verificó la resistencia del acero del anclaje en tensión, cortante e interacción."
            )
            docx_add_equation_line(doc, "N_sa = A_se,N · f_uta")
            docx_add_equation_line(doc, "V_sa = 0.60 · A_se,V · f_uta")
            docx_add_equation_line(doc, "(N_ua / φN_sa)^2 + (V_ua / φV_sa)^2 ≤ 1.0")
            docx_add_small_note(
                doc,
                "Referencia: ACI 318-25, Capítulo 17 (Anchor steel in tension and shear; interaction)."
            )
            docx_add_key_value_table(doc, [
                ("φNsa [kN]", f"{module5_results['phiNsa_kN']:.5f}"),
                ("φVsa [kN]", f"{module5_results['phiVsa_kN']:.5f}"),
                ("Nua por perno [kN]", f"{module5_results['Nua_per_bolt_kN']:.5f}"),
                ("Vua por perno [kN]", f"{module5_results['Vua_per_bolt_kN']:.5f}"),
                ("Interacción", f"{module5_results['interaction_value']:.5f}"),
                ("Chequeo", "Cumple" if module5_results["interaction_ok"] else "No cumple"),
            ])

        # Módulo 6 y 7
        if module6_results is not None:
            docx_add_heading(doc, "2.4 Concreto del anclaje en tensión", level=2)
            docx_add_normal_paragraph(
                doc,
                "Se evaluaron los modos de falla del concreto en tensión: concrete breakout, pullout y side-face blowout, "
                "adoptando la resistencia gobernante."
            )
            docx_add_equation_line(doc, "N_b = k_c λ_a √f'c · h_ef^1.5")
            docx_add_equation_line(doc, "A_Nco = 9 h_ef^2")
            docx_add_small_note(
                doc,
                "Referencia: ACI 318-25, Sección 17.6."
            )
            docx_add_key_value_table(doc, [
                ("φNcbg [kN]", f"{module6_results['phiNcbg_kN']:.5f}"),
                ("φNpn [kN]", "-" if module6_results["phiNpn_kN"] is None else f"{module6_results['phiNpn_kN']:.5f}"),
                ("φNsbg [kN]", "-" if module6_results["phiNsbg_kN"] is None else f"{module6_results['phiNsbg_kN']:.5f}"),
                ("Resistencia gobernante [kN]", f"{module6_results['phiNn_cg_kN']:.5f}"),
                ("Nua grupo [kN]", f"{module6_results['Nua_group_kN']:.5f}"),
                ("Chequeo", "Cumple" if module6_results["concrete_tension_ok"] else "No cumple"),
            ])

        if module7_results is not None:
            docx_add_heading(doc, "2.5 Concreto del anclaje en cortante", level=2)
            docx_add_normal_paragraph(
                doc,
                "Se evaluaron preliminarmente el concrete breakout en cortante, el pryout y la interacción N-V del concreto."
            )
            docx_add_equation_line(doc, "V_cpg = k_cp · N_cbg")
            docx_add_equation_line(doc, "(N_ua/φN_n)^(5/3) + (V_ua/φV_n)^(5/3) ≤ 1.0")
            docx_add_small_note(
                doc,
                "Referencia: ACI 318-25, Secciones 17.7 y 17.8."
            )
            docx_add_key_value_table(doc, [
                ("φVcbg [kN]", f"{module7_results['phiVcbg_kN']:.5f}"),
                ("φVcpg [kN]", f"{module7_results['phiVcpg_kN']:.5f}"),
                ("Resistencia gobernante [kN]", f"{module7_results['phiVn_cg_kN']:.5f}"),
                ("Vua grupo [kN]", f"{module7_results['Vua_group_kN']:.5f}"),
                ("Interacción concreta", f"{module7_results['interaction_concrete']:.5f}"),
                ("Chequeo", "Cumple" if module7_results["interaction_concrete_ok"] else "No cumple"),
            ])

        # Módulo 8
        if module8_results is not None:
            docx_add_heading(doc, "2.6 Requisitos geométricos mínimos", level=2)
            docx_add_normal_paragraph(
                doc,
                "Se verificaron espaciamiento mínimo, distancia al borde y límite de profundidad de empotramiento, "
                "según los requisitos geométricos mínimos de ACI para evitar splitting."
            )
            docx_add_small_note(
                doc,
                "Referencia: ACI 318-25, Sección 17.9."
            )
            docx_add_key_value_table(doc, [
                ("Espaciamiento real mínimo [mm]", f"{module8_results['min_spacing_real_mm']:.3f}"),
                ("Espaciamiento mínimo requerido [mm]", f"{module8_results['s_min_req_mm']:.3f}"),
                ("Distancia real mínima al borde [mm]", f"{module8_results['min_edge_real_mm']:.3f}"),
                ("Distancia mínima requerida [mm]", f"{module8_results['c_min_req_mm']:.3f}"),
                ("Chequeo global", "Cumple" if module8_results["geometric_ok"] else "No cumple"),
            ])

        # Módulo 9
        if module9_results is not None:
            docx_add_heading(doc, "2.7 Transferencia de cortante en la base", level=2)
            docx_add_normal_paragraph(
                doc,
                "Se evaluó el mecanismo resistente al cortante horizontal en la base, considerando fricción, "
                "anclajes y necesidad de shear key cuando corresponde."
            )
            docx_add_equation_line(doc, "V_n,fr = μ · P_u")
            docx_add_small_note(
                doc,
                "Referencia: AISC 360-22 J8/J9 y ACI 318-25 §17.11 cuando existe shear lug."
            )
            docx_add_key_value_table(doc, [
                ("Vu [kN]", f"{module9_results['Vu_kN']:.5f}"),
                ("Mecanismo", str(module9_results["selected_case"])),
                ("φVn considerado [kN]", f"{module9_results['phiVn_selected_kN']:.5f}"),
                ("¿Requiere shear key?", "Sí" if module9_results["shear_key_required"] else "No"),
                ("Vu remanente para key [kN]", f"{module9_results['Vu_remaining_for_key_kN']:.5f}"),
                ("Chequeo", "Cumple" if module9_results["shear_ok"] else "No cumple"),
            ])

        # Módulo 10
        if module10_results is not None:
            docx_add_heading(doc, "2.8 Soldadura", level=2)
            docx_add_normal_paragraph(
                doc,
                "Se revisó preliminarmente la soldadura de filete entre columna y placa base, y la soldadura del shear lug "
                "si este se requiere."
            )
            docx_add_equation_line(doc, "φR_n = φ · 0.60 · F_EXX · A_w")
            docx_add_equation_line(doc, "A_w = 0.707 · w · L_eff")
            docx_add_small_note(
                doc,
                "Referencia: AISC 360-22, Tabla J2.5 y Sección J2."
            )
            docx_add_key_value_table(doc, [
                ("Tamaño filete columna-base [mm]", f"{module10_results['column_weld_size_mm']:.3f}"),
                ("Longitud efectiva [mm]", f"{module10_results['L_col_weld_eff_mm']:.3f}"),
                ("φRn columna-base [kN]", f"{module10_results['phiRn_col_kN']:.5f}"),
                ("Demanda Tu [kN]", f"{module10_results['Tu_col_base_kN']:.5f}"),
                ("Chequeo columna-base", "Cumple" if module10_results["col_weld_ok"] else "No cumple"),
            ])

            if module10_results["shear_key_required"]:
                docx_add_key_value_table(doc, [
                    ("Vu lug [kN]", f"{module10_results['Vu_lug_kN']:.5f}"),
                    ("φRn lug [kN]", f"{module10_results['phiRn_lug_kN']:.5f}"),
                    ("Chequeo lug-placa", "Cumple" if module10_results["lug_weld_ok"] else "No cumple"),
                ], title="Soldadura shear lug - placa")

    elif analysis_mode == "Biaxial":
        docx_add_heading(doc, "2.1 Análisis biaxial", level=2)

        if module11_results is not None:
            docx_add_normal_paragraph(
                doc,
                "Se realizó primero una evaluación biaxial preliminar elástica, obteniendo presiones en las esquinas "
                "y una primera estimación de la tracción en pernos."
            )
            docx_add_equation_line(doc, "q(x,y) = P_u/A + (M_y/I_y)x + (M_x/I_x)y")
            docx_add_key_value_table(doc, [
                ("q_max [MPa]", f"{module11_results['q_max_MPa']:.5f}"),
                ("q_min [MPa]", f"{module11_results['q_min_MPa']:.5f}"),
                ("Compresión total", "Sí" if module11_results["full_compression"] else "No"),
                ("Tensión preliminar total [kN]", f"{module11_results['T_total_prelim_kN']:.5f}"),
                ("Perno crítico preliminar", "-" if module11_results["critical_bolt"] is None else str(module11_results["critical_bolt"])),
            ])

        if module12_results is not None:
            docx_add_normal_paragraph(
                doc,
                "Posteriormente se refinó la evaluación mediante integración por malla sobre la placa, "
                "separando zonas comprimidas y levantadas, y distribuyendo la tracción preliminar en el grupo de pernos."
            )
            docx_add_key_value_table(doc, [
                ("q_max [MPa]", f"{module12_results['q_max_MPa']:.5f}"),
                ("q_min [MPa]", f"{module12_results['q_min_MPa']:.5f}"),
                ("C total [kN]", f"{module12_results['C_kN']:.5f}"),
                ("T equivalente [kN]", f"{module12_results['T_equiv_kN']:.5f}"),
                ("xC [mm]", f"{module12_results['xC_mm']:.5f}"),
                ("yC [mm]", f"{module12_results['yC_mm']:.5f}"),
                ("Residual Mx [kN·m]", f"{module12_results['Mx_residual_kNm']:.5f}"),
                ("Residual My [kN·m]", f"{module12_results['My_residual_kNm']:.5f}"),
                ("Perno crítico refinado", "-" if module12_results["critical_bolt"] is None else str(module12_results["critical_bolt"])),
            ])

    # --------------------------------------------------------
    # 3. Gráficos
    # --------------------------------------------------------
    docx_add_heading(doc, "3. Gráficos relevantes", level=1)

    image_paths = []

    for key, fig in module14_figs.items():
        img_path = save_figure_to_temp(fig, f"report_{key}")
        image_paths.append((key, img_path))

    captions = {
        "uniaxial_pressure": "Figura 1. Distribución uniaxial de presión de contacto.",
        "uniaxial_anchor_tension": "Figura 2. Distribución de tracción en pernos para el análisis uniaxial.",
        "biaxial_elastic_pressure": "Figura 3. Mapa de presión biaxial elástica q(x,y).",
        "biaxial_contact_pressure": "Figura 4. Mapa de presión de contacto positiva q⁺(x,y).",
        "biaxial_anchor_tension": "Figura 5. Distribución refinada de tracción en pernos para el análisis biaxial.",
    }

    for key, path in image_paths:
        if os.path.exists(path):
            docx_add_image(doc, path, caption=captions.get(key, key), width_inches=5.8)

            interpretation = get_graph_interpretation(
                graph_key=key,
                analysis_mode=analysis_mode,
                module2_results=module2_results,
                module11_results=module11_results,
                module12_results=module12_results,
            )

            docx_add_normal_paragraph(doc, interpretation)
            doc.add_paragraph("")

    # --------------------------------------------------------
    # 4. Conclusión
    # --------------------------------------------------------
    docx_add_heading(doc, "4. Conclusión", level=1)

    status = module13_results["global_status"]

    if status == "Cumple":
        conclusion = (
            "Con base en las verificaciones realizadas, el diseño evaluado cumple los chequeos activos "
            "del modelo implementado."
        )
    elif status == "Preliminarmente aceptable":
        conclusion = (
            "El diseño resulta preliminarmente aceptable, aunque requiere validación adicional en los "
            "módulos todavía marcados como aproximados."
        )
    elif status == "Revisión requerida":
        conclusion = (
            "El diseño requiere revisión adicional. Existen chequeos que deben depurarse o complementarse "
            "antes de adoptar el detalle como solución final."
        )
    else:
        conclusion = (
            "El diseño no cumple uno o más chequeos relevantes y debe modificarse."
        )

    docx_add_normal_paragraph(doc, conclusion)

    if module13_results["critical_checks"]:
        docx_add_heading(doc, "4.1 Chequeos críticos", level=2)
        critical_rows = []
        for row in module13_results["critical_checks"]:
            critical_rows.append((
                f"{row['Módulo']} - {row['Chequeo']}",
                f"{row['Estado']} | {row['Nota']}"
            ))
        docx_add_key_value_table(doc, critical_rows)

    # --------------------------------------------------------
    # Guardar
    # --------------------------------------------------------
    file_path = os.path.join(tempfile.gettempdir(), "memoria_placa_base.docx")
    doc.save(file_path)

    return file_path
# ============================================================
# FUNCIONES DE GRÁFICO
# ============================================================

def _draw_dimension(ax, p1, p2, offset=20, text="", text_offset=5,
                    color="black", lw=1.2, fontsize=9):
    x1, y1 = p1
    x2, y2 = p2

    if abs(y2 - y1) < 1e-9:
        y_dim = y1 + offset
        ax.plot([x1, x1], [y1, y_dim], color=color, linewidth=lw)
        ax.plot([x2, x2], [y2, y_dim], color=color, linewidth=lw)
        ax.annotate("", xy=(x1, y_dim), xytext=(x2, y_dim),
                    arrowprops=dict(arrowstyle="<->", color=color, linewidth=lw))
        ax.text((x1 + x2) / 2, y_dim + text_offset, text,
                ha="center", va="bottom", color=color, fontsize=fontsize)

    elif abs(x2 - x1) < 1e-9:
        x_dim = x1 + offset
        ax.plot([x1, x_dim], [y1, y1], color=color, linewidth=lw)
        ax.plot([x2, x_dim], [y2, y2], color=color, linewidth=lw)
        ax.annotate("", xy=(x_dim, y1), xytext=(x_dim, y2),
                    arrowprops=dict(arrowstyle="<->", color=color, linewidth=lw))
        ax.text(x_dim + text_offset, (y1 + y2) / 2, text,
                ha="left", va="center", color=color, fontsize=fontsize, rotation=90)


def _draw_steel_section(ax, column_plot: dict):
    section_type = column_plot.get("section_type", "W").upper()
    d = column_plot["d_mm"]
    bf = column_plot["bf_mm"]

    if section_type in ["W", "H", "I"]:
        tf = column_plot["tf_mm"]
        tw = column_plot["tw_mm"]

        top_flange = Rectangle((-bf / 2, d / 2 - tf), bf, tf,
                               fill=False, linewidth=2.0, linestyle="--", edgecolor="gray")
        bottom_flange = Rectangle((-bf / 2, -d / 2), bf, tf,
                                  fill=False, linewidth=2.0, linestyle="--", edgecolor="gray")
        web = Rectangle((-tw / 2, -d / 2 + tf), tw, d - 2 * tf,
                        fill=False, linewidth=2.0, linestyle="--", edgecolor="gray")

        ax.add_patch(top_flange)
        ax.add_patch(bottom_flange)
        ax.add_patch(web)

    else:
        rect = Rectangle((-bf / 2, -d / 2), bf, d,
                         fill=False, linewidth=2.0, linestyle="--", edgecolor="gray")
        ax.add_patch(rect)


def base_plate_layout_plot(
    base_plate_plot: dict,
    pedestal_plot: dict,
    column_plot: dict,
    bolt_df: pd.DataFrame,
    anchors: AnchorLayout,
):
    fig, ax = plt.subplots(figsize=(10, 10))

    Bp = pedestal_plot["B_mm"]
    Np = pedestal_plot["N_mm"]
    B = base_plate_plot["B_mm"]
    N = base_plate_plot["N_mm"]
    d = column_plot["d_mm"]
    bf = column_plot["bf_mm"]

    edge_x = anchors.edge_x_mm
    edge_y = anchors.edge_y_mm
    r = anchors.db_mm / 2.0

    x_ped_l = -Bp / 2
    x_ped_r = Bp / 2
    y_ped_b = -Np / 2
    y_ped_t = Np / 2

    x_pla_l = -B / 2
    x_pla_r = B / 2
    y_pla_b = -N / 2
    y_pla_t = N / 2

    x_col_l = -bf / 2
    x_col_r = bf / 2
    y_col_b = -d / 2
    y_col_t = d / 2

    x_bolt_left = x_pla_l + edge_x
    x_bolt_right = x_pla_r - edge_x
    y_bolt_bottom = y_pla_b + edge_y
    y_bolt_top = y_pla_t - edge_y

    ax.add_patch(Rectangle((x_ped_l, y_ped_b), Bp, Np,
                           fill=False, linewidth=2.0, edgecolor="saddlebrown"))
    ax.add_patch(Rectangle((x_pla_l, y_pla_b), B, N,
                           fill=False, linewidth=2.4, edgecolor="black"))

    _draw_steel_section(ax, column_plot)

    for _, row in bolt_df.iterrows():
        x = row["x_mm"]
        y = row["y_mm"]
        ax.add_patch(Circle((x, y), radius=r, fill=False, linewidth=1.8, edgecolor="blue"))
        ax.plot(x, y, "bo", markersize=3)
        ax.text(x + 6, y + 6, str(int(row["Perno"])), fontsize=9, color="blue")

    ax.plot(0, 0, marker="+", markersize=14, markeredgewidth=2.2, color="red")
    ax.axhline(0, linestyle=":", color="red", linewidth=1.0)
    ax.axvline(0, linestyle=":", color="red", linewidth=1.0)

    _draw_dimension(ax, (x_ped_l, y_ped_b), (x_ped_r, y_ped_b),
                    offset=-55, text=f"Pedestal B = {Bp:.0f} mm", color="saddlebrown")
    _draw_dimension(ax, (x_ped_l, y_ped_b), (x_ped_l, y_ped_t),
                    offset=-55, text=f"Pedestal N = {Np:.0f} mm", color="saddlebrown")
    _draw_dimension(ax, (x_pla_l, y_pla_t), (x_pla_r, y_pla_t),
                    offset=35, text=f"Placa B = {B:.0f} mm", color="black")
    _draw_dimension(ax, (x_pla_r, y_pla_b), (x_pla_r, y_pla_t),
                    offset=35, text=f"Placa N = {N:.0f} mm", color="black")
    _draw_dimension(ax, (x_col_l, y_col_t), (x_col_r, y_col_t),
                    offset=18, text=f"bf = {bf:.0f} mm", color="gray")
    _draw_dimension(ax, (x_col_r, y_col_b), (x_col_r, y_col_t),
                    offset=18, text=f"d = {d:.0f} mm", color="gray")
    _draw_dimension(ax, (x_pla_l, y_bolt_bottom), (x_bolt_left, y_bolt_bottom),
                    offset=-35, text=f"ex = {edge_x:.0f} mm", color="blue")
    _draw_dimension(ax, (x_bolt_left, y_pla_b), (x_bolt_left, y_bolt_bottom),
                    offset=-35, text=f"ey = {edge_y:.0f} mm", color="blue")

    margin = max(Bp, Np) * 0.22
    ax.set_xlim(-Bp / 2 - margin, Bp / 2 + margin)
    ax.set_ylim(-Np / 2 - margin, Np / 2 + margin)
    ax.set_aspect("equal", adjustable="box")
    ax.grid(True, linestyle="--", alpha=0.35)
    ax.set_xlabel("x [mm]")
    ax.set_ylabel("y [mm]")
    ax.set_title("Geometría de placa base")

    return fig


# ============================================================
# SIDEBAR - ENTRADA DE DATOS
# ============================================================

st.sidebar.header("Entrada de datos")

with st.sidebar.expander("Cargas", expanded=True):
    Pu_kN = st.number_input("Pu [kN]", min_value=0.001, value=1200.0, step=10.0)
    Mux_kNm = st.number_input("Mux [kN·m]", value=80.0, step=5.0)
    Muy_kNm = st.number_input("Muy [kN·m]", value=55.0, step=5.0)
    Vux_kN = st.number_input("Vux [kN]", value=40.0, step=5.0)
    Vuy_kN = st.number_input("Vuy [kN]", value=20.0, step=5.0)

with st.sidebar.expander("Materiales", expanded=True):
    Fy_plate_MPa = st.number_input("Fy placa [MPa]", min_value=0.001, value=250.0)
    Fu_plate_MPa = st.number_input("Fu placa [MPa]", min_value=0.001, value=400.0)
    fc_MPa = st.number_input("f'c [MPa]", min_value=0.001, value=28.0)
    Fy_anchor_MPa = st.number_input("Fy anclaje [MPa]", min_value=0.001, value=414.0)
    Fu_anchor_MPa = st.number_input("Fu anclaje [MPa]", min_value=0.001, value=620.0)

with st.sidebar.expander("Columna", expanded=True):
    section_type = st.selectbox("Tipo de perfil", ["W", "H", "I", "HSS", "BOX"], index=0)
    d_mm = st.number_input("d [mm]", min_value=0.001, value=300.0)
    bf_mm = st.number_input("bf [mm]", min_value=0.001, value=250.0)
    tf_mm = st.number_input("tf [mm]", min_value=0.0, value=14.0)
    tw_mm = st.number_input("tw [mm]", min_value=0.0, value=9.0)

with st.sidebar.expander("Placa base", expanded=True):
    B_bp_mm = st.number_input("B placa [mm]", min_value=0.001, value=600.0)
    N_bp_mm = st.number_input("N placa [mm]", min_value=0.001, value=600.0)
    tp_mm = st.number_input("tp [mm]", min_value=0.001, value=25.0)

with st.sidebar.expander("Pernos", expanded=True):
    nbx = st.number_input("nbx", min_value=2, value=4, step=1)
    nby = st.number_input("nby", min_value=2, value=5, step=1)
    edge_x_mm = st.number_input("edge_x [mm]", min_value=0.001, value=100.0)
    edge_y_mm = st.number_input("edge_y [mm]", min_value=0.001, value=100.0)
    db_mm = st.number_input("db [mm]", min_value=0.001, value=20.0)
    Ab_mm2 = st.number_input("Ab [mm²]", min_value=0.001, value=245.0)
    hef_mm = st.number_input("hef [mm]", min_value=0.001, value=300.0)

with st.sidebar.expander("Módulo 5 - acero del anclaje", expanded=False):
    anchor_type = st.selectbox(
        "Tipo de anclaje para acero",
        ["headed_bolt", "hooked_bolt", "headed_stud", "adhesive_anchor"],
        index=0
    )

    use_built_up_grout_pad = st.checkbox(
        "¿Hay built-up grout pad?",
        value=False
    )

    phi_anchor_tension_steel = st.number_input(
        "φ acero en tensión",
        min_value=0.01,
        max_value=1.00,
        value=0.75,
        step=0.01
    )

    phi_anchor_shear_steel = st.number_input(
        "φ acero en cortante",
        min_value=0.01,
        max_value=1.00,
        value=0.65,
        step=0.01
    )

with st.sidebar.expander("Módulo 6 - concreto en tensión", expanded=False):
    anchor_installation = st.selectbox(
        "Instalación del anclaje",
        ["cast_in", "post_installed"],
        index=0
    )

with st.sidebar.expander("Módulo 7 - concreto en cortante", expanded=False):
    phi_concrete_shear = st.number_input(
        "φ concreto en cortante",
        min_value=0.01,
        max_value=1.00,
        value=0.70,
        step=0.01
    )
with st.sidebar.expander("Módulo 8 - ACI 17.9 geometría mínima", expanded=False):
    supplementary_reinforcement_for_splitting = st.checkbox(
        "¿Existe refuerzo suplementario para controlar splitting?",
        value=False
    )

    anchor_torqued = st.checkbox(
        "¿El anclaje será torqued?",
        value=True
    )

    nominal_max_agg_mm = st.number_input(
        "Tamaño máximo nominal del agregado [mm]",
        min_value=1.0,
        value=19.0,
        step=1.0
    )

    required_cover_mm = st.number_input(
        "Recubrimiento mínimo requerido según 20.5.1.3 [mm]",
        min_value=0.0,
        value=75.0,
        step=5.0
    )

    product_specific_geometry_data = st.checkbox(
        "¿Existe información product-specific ACI 355.2 / 355.4?",
        value=False
    )

    product_specific_min_edge_mm = st.number_input(
        "Distancia mínima al borde product-specific [mm]",
        min_value=0.0,
        value=0.0,
        step=5.0
    )

    product_specific_min_spacing_mm = st.number_input(
        "Espaciamiento mínimo product-specific [mm]",
        min_value=0.0,
        value=0.0,
        step=5.0
    )

    post_installed_type_for_17_9 = st.selectbox(
        "Tipo de post-installed para 17.9",
        ["adhesive", "torque_controlled", "displacement_controlled", "screw", "undercut"],
        index=0
    )

    tests_permit_greater_hef = st.checkbox(
        "¿Existen ensayos que permiten exceder el límite de hef de 17.9.4?",
        value=False
    )
    ca1_shear_mm = st.number_input(
        "ca1, shear [mm] (borde en dirección del cortante)",
        min_value=0.001,
        value=250.0,
        step=5.0
    )

    member_thickness_for_shear_mm = st.number_input(
        "ha [mm] (espesor del miembro para shear breakout)",
        min_value=0.001,
        value=500.0,
        step=5.0
    )
    service_cracked = st.checkbox(
        "¿La región está agrietada a nivel de servicio?",
        value=True
    )

    lambda_a = st.number_input(
        "λa",
        min_value=0.10,
        max_value=2.00,
        value=1.00,
        step=0.05
    )

    psi_a = st.number_input(
        "ψa",
        min_value=0.10,
        max_value=2.00,
        value=1.00,
        step=0.05
    )

    phi_concrete_tension = st.number_input(
        "φ concreto en tensión",
        min_value=0.01,
        max_value=1.00,
        value=0.70,
        step=0.01
    )

    # Para headed bolts / headed studs
    Abrg_mm2 = st.number_input(
        "Abrg [mm²] (área neta de apoyo de cabeza)",
        min_value=0.001,
        value=800.0,
        step=10.0
    )

    # Para hooked bolts
    eh_mm = st.number_input(
        "eh [mm] (solo J/L bolt)",
        min_value=0.0,
        value=0.0,
        step=5.0
    )

    # Distancias a borde del grupo respecto al pedestal
    ca1_x_mm = st.number_input(
        "ca1,x [mm] (borde mínimo en x)",
        min_value=0.001,
        value=250.0,
        step=5.0
    )

    ca1_y_mm = st.number_input(
        "ca1,y [mm] (borde mínimo en y)",
        min_value=0.001,
        value=250.0,
        step=5.0
    )

    ca2_mm = st.number_input(
        "ca2 [mm] (borde perpendicular para side-face blowout)",
        min_value=0.001,
        value=300.0,
        step=5.0
    )

with st.sidebar.expander("Módulo 9 - cortante en la base", expanded=False):
    base_shear_mode = st.selectbox(
        "Modo de evaluación del cortante",
        ["auto_detect", "manual"],
        index=0
    )

    base_shear_mechanism = st.selectbox(
        "Mecanismo manual (si eliges manual)",
        ["friction", "anchors", "shear_key", "combined"],
        index=0
    )

    mu = st.number_input(
        "μ coeficiente de fricción",
        min_value=0.0,
        max_value=2.0,
        value=0.30,
        step=0.01
    )

    phi_base_friction = st.number_input(
        "φ fricción en la base",
        min_value=0.01,
        max_value=1.00,
        value=1.00,
        step=0.01
    )

    phiVn_shear_key_external_kN = st.number_input(
        "φVn de shear key / shear lug [kN] (si ya está definido)",
        min_value=0.0,
        value=0.0,
        step=5.0
    )

    allow_combined_mechanisms = st.checkbox(
        "¿Permitir combinación fricción + anclajes?",
        value=False
    )

with st.sidebar.expander("Módulo 10 - soldadura", expanded=False):
    FEXX_MPa = st.number_input(
        "FEXX del electrodo [MPa]",
        min_value=1.0,
        value=490.0,
        step=10.0
    )

    phi_weld = st.number_input(
        "φ soldadura",
        min_value=0.01,
        max_value=1.00,
        value=0.75,
        step=0.01
    )

    # ----------------------------------------------------
    # Soldadura columna - placa
    # ----------------------------------------------------
    column_weld_size_mm = st.number_input(
        "Tamaño filete columna-placa [mm]",
        min_value=0.0,
        value=8.0,
        step=1.0
    )

    column_weld_layout = st.selectbox(
        "Disposición de soldadura columna-placa",
        [
            "auto_by_section",
            "flanges_only",
            "web_only",
            "flanges_plus_web",
            "all_perimeter_rect",
            "manual"
        ],
        index=0
    )

    column_weld_length_manual_mm = st.number_input(
        "Longitud efectiva manual columna-placa [mm]",
        min_value=0.0,
        value=0.0,
        step=10.0
    )

    # ----------------------------------------------------
    # Soldadura shear lug - placa
    # ----------------------------------------------------
    provide_shear_lug_weld = st.checkbox(
        "¿Evaluar soldadura shear lug-placa?",
        value=False
    )

    shear_lug_weld_size_mm = st.number_input(
        "Tamaño filete shear lug-placa [mm]",
        min_value=0.0,
        value=8.0,
        step=1.0
    )

    shear_lug_weld_length_mm = st.number_input(
        "Longitud efectiva total shear lug-placa [mm]",
        min_value=0.0,
        value=200.0,
        step=10.0
    )

with st.sidebar.expander("Módulo 12 - refinamiento biaxial", expanded=False):
    biaxial_grid_nx = st.number_input(
        "Número de divisiones en x",
        min_value=11,
        max_value=201,
        value=41,
        step=2
    )

    biaxial_grid_ny = st.number_input(
        "Número de divisiones en y",
        min_value=11,
        max_value=201,
        value=41,
        step=2
    )

    biaxial_use_bolt_uplift_distribution = st.checkbox(
        "Distribuir tracción biaxial preliminar en pernos",
        value=True
    )

with st.sidebar.expander("Módulo 14 - gráficas", expanded=False):
    show_elastic_biaxial_pressure = st.checkbox(
        "Mostrar presión elástica biaxial q(x,y)",
        value=True
    )

    show_contact_biaxial_pressure = st.checkbox(
        "Mostrar presión de contacto q⁺(x,y)",
        value=True
    )

    uniaxial_plot_points = st.number_input(
        "Puntos para gráfica uniaxial",
        min_value=21,
        max_value=1001,
        value=201,
        step=20
    )    
with st.sidebar.expander("Pedestal", expanded=True):
    B_ped_mm = st.number_input("B pedestal [mm]", min_value=0.001, value=900.0)
    N_ped_mm = st.number_input("N pedestal [mm]", min_value=0.001, value=900.0)
    h_ped_mm = st.number_input("h pedestal [mm]", min_value=0.001, value=500.0)

with st.sidebar.expander("Modo de análisis", expanded=True):
    analysis_mode = st.selectbox(
        "Selecciona el modo",
        ["Uniaxial", "Biaxial"],
        index=0
    )
    uniaxial_axis = st.selectbox(
        "Eje uniaxial",
        ["x", "y"],
        index=0
    )


# ============================================================
# ARMAR OBJETOS
# ============================================================

loads = Loads(
    Pu_kN=Pu_kN,
    Mux_kNm=Mux_kNm,
    Muy_kNm=Muy_kNm,
    Vux_kN=Vux_kN,
    Vuy_kN=Vuy_kN,
)

materials = Materials(
    Fy_plate_MPa=Fy_plate_MPa,
    Fu_plate_MPa=Fu_plate_MPa,
    fc_MPa=fc_MPa,
    Fy_anchor_MPa=Fy_anchor_MPa,
    Fu_anchor_MPa=Fu_anchor_MPa,
)

column = ColumnGeometry(
    d_col_mm=d_mm,
    bf_col_mm=bf_mm,
)

column_plot = {
    "section_type": section_type,
    "d_mm": d_mm,
    "bf_mm": bf_mm,
    "tf_mm": tf_mm,
    "tw_mm": tw_mm,
}

base_plate = BasePlateGeometry(
    B_bp_mm=B_bp_mm,
    N_bp_mm=N_bp_mm,
    tp_mm=tp_mm,
)

anchors = AnchorLayout(
    nbx=int(nbx),
    nby=int(nby),
    edge_x_mm=edge_x_mm,
    edge_y_mm=edge_y_mm,
    db_mm=db_mm,
    Ab_mm2=Ab_mm2,
    hef_mm=hef_mm,
)

pedestal = PedestalGeometry(
    B_ped_mm=B_ped_mm,
    N_ped_mm=N_ped_mm,
    h_ped_mm=h_ped_mm,
)


# ============================================================
# EJECUCIÓN BASE EN CADENA
# ============================================================

try:
    validate_inputs(loads, materials, column, base_plate, anchors, pedestal)

    bolt_df = generate_bolt_coordinates(
        B_bp_mm=base_plate.B_bp_mm,
        N_bp_mm=base_plate.N_bp_mm,
        nbx=anchors.nbx,
        nby=anchors.nby,
        edge_x_mm=anchors.edge_x_mm,
        edge_y_mm=anchors.edge_y_mm,
    )

    summary_df = build_input_summary(
        loads=loads,
        materials=materials,
        column=column,
        base_plate=base_plate,
        anchors=anchors,
        pedestal=pedestal,
    )

    layout_info = compute_layout_parameters(base_plate, anchors)

    base_plate_plot = {
        "B_mm": base_plate.B_bp_mm,
        "N_mm": base_plate.N_bp_mm,
        "tp_mm": base_plate.tp_mm,
    }

    pedestal_plot = {
        "B_mm": pedestal.B_ped_mm,
        "N_mm": pedestal.N_ped_mm,
        "h_mm": pedestal.h_ped_mm,
    }

    # --------------------------------------------------------
    # MÓDULO 1
    # --------------------------------------------------------
    if analysis_mode == "Uniaxial":
        module1_results = module1_uniaxial_preliminary(
            loads=loads,
            base_plate=base_plate,
            axis=uniaxial_axis,
        )
    else:
        module1_results = None
    # --------------------------------------------------------
    # MÓDULO 2
    # --------------------------------------------------------
    if analysis_mode == "Uniaxial":
        module2_results = module2_uniaxial_bearing(
            loads=loads,
            base_plate=base_plate,
            materials=materials,
            axis=uniaxial_axis,
        )
    else:
        module2_results = None

    # --------------------------------------------------------
    # MÓDULO 3
    # --------------------------------------------------------
    if analysis_mode == "Uniaxial":
        module3_results = module3_uniaxial_anchor_tension(
            loads=loads,
            base_plate=base_plate,
            anchors=anchors,
            bolt_df=bolt_df,
            module2_results=module2_results,
        )
    else:
        module3_results = None   

    # --------------------------------------------------------
    # MÓDULO 4
    # --------------------------------------------------------
    if analysis_mode == "Uniaxial":
        module4_results = module4_plate_thickness(
            base_plate=base_plate,
            column=column,
            materials=materials,
            module2_results=module2_results,
        )
    else:
        module4_results = None

    # --------------------------------------------------------
    # MÓDULO 5
    # --------------------------------------------------------
    if analysis_mode == "Uniaxial":
        module5_results = module5_anchor_steel_strength(
            loads=loads,
            materials=materials,
            anchors=anchors,
            layout_info=layout_info,
            module3_results=module3_results,
            analysis_mode=analysis_mode,
            uniaxial_axis=uniaxial_axis,
            anchor_type=anchor_type,
            phi_anchor_tension_steel=phi_anchor_tension_steel,
            phi_anchor_shear_steel=phi_anchor_shear_steel,
            use_built_up_grout_pad=use_built_up_grout_pad,
        )
    else:
        module5_results = None

    # --------------------------------------------------------
    # MÓDULO 6
    # --------------------------------------------------------
    if analysis_mode == "Uniaxial":
        module6_results = module6_concrete_tension(
            materials=materials,
            anchors=anchors,
            pedestal=pedestal,
            bolt_df=bolt_df,
            module3_results=module3_results,
            anchor_type=anchor_type,
            anchor_installation=anchor_installation,
            service_cracked=service_cracked,
            lambda_a=lambda_a,
            psi_a=psi_a,
            phi_concrete_tension=phi_concrete_tension,
            Abrg_mm2=Abrg_mm2,
            eh_mm=eh_mm,
            ca1_x_mm=ca1_x_mm,
            ca1_y_mm=ca1_y_mm,
            ca2_mm=ca2_mm,
        )
    else:
        module6_results = None

    # --------------------------------------------------------
    # MÓDULO 7
    # --------------------------------------------------------
    if analysis_mode == "Uniaxial":
        module7_results = module7_concrete_shear(
            loads=loads,
            anchors=anchors,
            bolt_df=bolt_df,
            module6_results=module6_results,
            analysis_mode=analysis_mode,
            uniaxial_axis=uniaxial_axis,
            lambda_a=lambda_a,
            phi_concrete_shear=phi_concrete_shear,
            ca1_shear_mm=ca1_shear_mm,
            member_thickness_for_shear_mm=member_thickness_for_shear_mm,
        )
    else:
        module7_results = None

    # --------------------------------------------------------
    # MÓDULO 8
    # --------------------------------------------------------
    if analysis_mode == "Uniaxial":
        module8_results = module8_geometry_minimums_aci_17_9(
            anchors=anchors,
            pedestal=pedestal,
            bolt_df=bolt_df,
            anchor_installation=anchor_installation,
            post_installed_type_for_17_9=post_installed_type_for_17_9,
            anchor_torqued=anchor_torqued,
            supplementary_reinforcement_for_splitting=supplementary_reinforcement_for_splitting,
            nominal_max_agg_mm=nominal_max_agg_mm,
            required_cover_mm=required_cover_mm,
            product_specific_geometry_data=product_specific_geometry_data,
            product_specific_min_edge_mm=product_specific_min_edge_mm,
            product_specific_min_spacing_mm=product_specific_min_spacing_mm,
            tests_permit_greater_hef=tests_permit_greater_hef,
        )
    else:
        module8_results = None   

    # --------------------------------------------------------
    # MÓDULO 9
    # --------------------------------------------------------
    if analysis_mode == "Uniaxial":
        module9_results = module9_base_shear_transfer(
            loads=loads,
            module5_results=module5_results,
            analysis_mode=analysis_mode,
            uniaxial_axis=uniaxial_axis,
            base_shear_mode=base_shear_mode,
            base_shear_mechanism=base_shear_mechanism,
            mu=mu,
            phi_base_friction=phi_base_friction,
            phiVn_shear_key_external_kN=phiVn_shear_key_external_kN,
            allow_combined_mechanisms=allow_combined_mechanisms,
        )
    else:
        module9_results = None    

    # --------------------------------------------------------
    # MÓDULO 10
    # --------------------------------------------------------
    if analysis_mode == "Uniaxial":
        module10_results = module10_weld_design(
            column_plot=column_plot,
            loads=loads,
            module3_results=module3_results,
            module9_results=module9_results,
            FEXX_MPa=FEXX_MPa,
            phi_weld=phi_weld,
            column_weld_size_mm=column_weld_size_mm,
            column_weld_layout=column_weld_layout,
            column_weld_length_manual_mm=column_weld_length_manual_mm,
            provide_shear_lug_weld=provide_shear_lug_weld,
            shear_lug_weld_size_mm=shear_lug_weld_size_mm,
            shear_lug_weld_length_mm=shear_lug_weld_length_mm,
        )
    else:
        module10_results = None
    # --------------------------------------------------------
    # MÓDULO 11
    # --------------------------------------------------------
    if analysis_mode == "Biaxial":
        module11_results = module11_biaxial_preliminary(
            loads=loads,
            base_plate=base_plate,
            bolt_df=bolt_df,
        )
    else:
        module11_results = None
    # --------------------------------------------------------
    # MÓDULO 12
    # --------------------------------------------------------
    if analysis_mode == "Biaxial":
        module12_results = module12_biaxial_grid_refinement(
            loads=loads,
            base_plate=base_plate,
            bolt_df=bolt_df,
            nx=biaxial_grid_nx,
            ny=biaxial_grid_ny,
            use_bolt_distribution=biaxial_use_bolt_uplift_distribution,
        )
    else:
        module12_results = None

    # --------------------------------------------------------
    # MÓDULO 13
    # --------------------------------------------------------
    module13_results = module13_design_summary(
        analysis_mode=analysis_mode,
        module2_results=module2_results if analysis_mode == "Uniaxial" else None,
        module4_results=module4_results if analysis_mode == "Uniaxial" else None,
        module5_results=module5_results if analysis_mode == "Uniaxial" else None,
        module6_results=module6_results if analysis_mode == "Uniaxial" else None,
        module7_results=module7_results if analysis_mode == "Uniaxial" else None,
        module8_results=module8_results if analysis_mode == "Uniaxial" else None,
        module9_results=module9_results if analysis_mode == "Uniaxial" else None,
        module10_results=module10_results if analysis_mode == "Uniaxial" else None,
        module11_results=module11_results if analysis_mode == "Biaxial" else None,
        module12_results=module12_results if analysis_mode == "Biaxial" else None,
    )
    # --------------------------------------------------------
    # MÓDULO 14 - GRÁFICAS
    # --------------------------------------------------------
    module14_figs = {}

    if analysis_mode == "Uniaxial":
        module14_figs["uniaxial_pressure"] = plot_uniaxial_pressure(
            module2_results=module2_results,
            base_plate=base_plate,
            axis=uniaxial_axis,
            npts=uniaxial_plot_points,
        )

        module14_figs["uniaxial_anchor_tension"] = plot_uniaxial_anchor_tension(
            base_plate=base_plate,
            bolt_df=bolt_df,
            module3_results=module3_results,
            axis=uniaxial_axis,
        )

    elif analysis_mode == "Biaxial":
        if show_elastic_biaxial_pressure:
            module14_figs["biaxial_elastic_pressure"] = plot_biaxial_pressure_field(
                module12_results=module12_results,
                title="Presión biaxial elástica q(x,y)",
                use_positive_only=False,
            )

        if show_contact_biaxial_pressure:
            module14_figs["biaxial_contact_pressure"] = plot_biaxial_pressure_field(
                module12_results=module12_results,
                title="Presión de contacto q⁺(x,y)",
                use_positive_only=True,
            )

        module14_figs["biaxial_anchor_tension"] = plot_biaxial_anchor_tension(
            base_plate=base_plate,
            module12_results=module12_results,
        )
    # --------------------------------------------------------
    # PESTAÑAS
    # --------------------------------------------------------
    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9, tab10, tab11, tab12, tab13, tab14, tab15, tab16, tab17 = st.tabs([
    "Datos y geometría",
    "Módulo 1 - Uniaxial",
    "Módulo 2 - Compresión",
    "Módulo 3 - Pernos",
    "Módulo 4 - Espesor",
    "Módulo 5 - Acero anclaje",
    "Módulo 6 - Concreto tensión",
    "Módulo 7 - Concreto cortante",
    "Módulo 8 - ACI 17.9",
    "Módulo 9 - Cortante base",
    "Módulo 10 - Soldadura",
    "Módulo 11 - Biaxial preliminar",
    "Módulo 12 - Biaxial refinado",
    "Módulo 13 - Cierre",
    "Módulo 14 - Gráficas",
    "Resumen",
    "Estado",
    ])

    with tab1:
        st.subheader("Resumen de datos de entrada")
        st.dataframe(summary_df, use_container_width=True, hide_index=True)

        st.subheader("Parámetros geométricos")
        layout_df = pd.DataFrame(
            {"Parámetro": list(layout_info.keys()), "Valor": list(layout_info.values())}
        )
        st.dataframe(layout_df, use_container_width=True, hide_index=True)

        st.subheader("Coordenadas de pernos")
        st.dataframe(bolt_df, use_container_width=True, hide_index=True)

        st.subheader("Gráfico geométrico")
        fig = base_plate_layout_plot(
            base_plate_plot=base_plate_plot,
            pedestal_plot=pedestal_plot,
            column_plot=column_plot,
            bolt_df=bolt_df,
            anchors=anchors,
        )
        st.pyplot(fig, clear_figure=True)

    with tab2:
        st.subheader("Módulo 1 - análisis uniaxial preliminar")

        if analysis_mode != "Uniaxial":
            st.info("En la barra lateral elegiste modo Biaxial. Cambia a Uniaxial para activar este módulo.")
        else:
            mod1_df = pd.DataFrame({
                "Parámetro": [
                    "Eje analizado",
                    "Pu [kN]",
                    "Mu [kN·m]",
                    "Área de placa [mm²]",
                    "Presión media q [MPa]",
                    "Excentricidad e [mm]",
                    "Núcleo central (dim/6) [mm]",
                    "Relación e/(dim/6)",
                    "Compresión total preliminar",
                    "Posible levantamiento preliminar",
                ],
                "Valor": [
                    module1_results["axis"],
                    f"{module1_results['Pu_kN']:.3f}",
                    f"{module1_results['Mu_kNm']:.3f}",
                    f"{module1_results['A_plate_mm2']:.3f}",
                    f"{module1_results['q_avg_MPa']:.5f}",
                    f"{module1_results['e_mm']:.3f}",
                    f"{module1_results['kern_mm']:.3f}",
                    f"{module1_results['e_over_kern']:.3f}",
                    "Sí" if module1_results["full_compression"] else "No",
                    "Sí" if module1_results["possible_uplift"] else "No",
                ],
            })
            st.dataframe(mod1_df, use_container_width=True, hide_index=True)

            if module1_results["possible_uplift"]:
                st.warning(
                    "La excentricidad supera el núcleo central en este análisis preliminar. "
                    "En los siguientes módulos deberá analizarse compresión parcial y tracción en pernos."
                )
            else:
                st.success(
                    "La resultante cae dentro del núcleo central en este análisis preliminar."
                )

    with tab3:
        st.subheader("Módulo 2 - compresión uniaxial bajo placa")

        if analysis_mode != "Uniaxial":
            st.info("En la barra lateral elegiste modo Biaxial. Cambia a Uniaxial para activar este módulo.")
        else:
            mod2_df = pd.DataFrame({
                "Parámetro": [
                    "Eje analizado",
                    "Pu [kN]",
                    "Mu [kN·m]",
                    "Dimensión resistente L [mm]",
                    "Ancho transversal b [mm]",
                    "Área total de placa [mm²]",
                    "Excentricidad e [mm]",
                    "Núcleo central L/6 [mm]",
                    "Caso de compresión",
                    "Presión media q_prom [MPa]",
                    "Presión máxima q_max [MPa]",
                    "Presión mínima q_min [MPa]",
                    "Longitud comprimida a [mm]",
                    "Resultante C desde centro [mm]",
                    "φ bearing",
                    "φ·0.85·f'c [MPa]",
                    "Chequeo preliminar de bearing",
                ],
                "Valor": [
                    module2_results["axis"],
                    f"{module2_results['Pu_kN']:.3f}",
                    f"{module2_results['Mu_kNm']:.3f}",
                    f"{module2_results['L_mm']:.3f}",
                    f"{module2_results['b_mm']:.3f}",
                    f"{module2_results['A_mm2']:.3f}",
                    f"{module2_results['e_mm']:.3f}",
                    f"{module2_results['kern_mm']:.3f}",
                    "Compresión total" if module2_results["case"] == "full_compression" else "Compresión parcial",
                    f"{module2_results['q_avg_MPa']:.5f}",
                    f"{module2_results['q_max_MPa']:.5f}",
                    f"{module2_results['q_min_MPa']:.5f}",
                    f"{module2_results['a_comp_mm']:.3f}",
                    f"{module2_results['xC_from_center_mm']:.3f}",
                    f"{module2_results['phi_bearing']:.3f}",
                    f"{module2_results['q_allow_phi_MPa']:.5f}",
                    "Cumple" if module2_results["bearing_ok"] else "No cumple",
                ],
            })
            st.dataframe(mod2_df, use_container_width=True, hide_index=True)

            if module2_results["case"] == "full_compression":
                st.success(
                    "La placa permanece totalmente comprimida en el análisis uniaxial simplificado."
                )
            else:
                st.warning(
                    "Se detecta compresión parcial en el análisis uniaxial simplificado. "
                    "En el siguiente módulo se deberá calcular la tracción en pernos."
                )

            if module2_results["bearing_ok"]:
                st.success(
                    "El chequeo preliminar de bearing con q_max ≤ φ·0.85·f'c cumple."
                )
            else:
                st.error(
                    "El chequeo preliminar de bearing con q_max ≤ φ·0.85·f'c no cumple."
                )    

    
    with tab4:
        st.subheader("Módulo 3 - tracción uniaxial en pernos")

        if analysis_mode != "Uniaxial":
            st.info("En la barra lateral elegiste modo Biaxial. Cambia a Uniaxial para activar este módulo.")
        else:
            mod3_df = pd.DataFrame({
                "Parámetro": [
                    "Eje analizado",
                    "Caso de compresión",
                    "Tracción activa en pernos",
                    "Lado traccionado",
                    "Brazo interno C-T [mm]",
                    "Número de pernos traccionados",
                    "Tracción total T [kN]",
                    "Tracción por perno [kN]",
                    "Fila extrema de pernos [coord. mm]",
                    "Pernos críticos",
                ],
                "Valor": [
                    module3_results["axis"],
                    "Compresión total" if module3_results["case"] == "full_compression" else "Compresión parcial",
                    "Sí" if module3_results["tension_active"] else "No",
                    module3_results["tension_side"],
                    f"{module3_results['lever_arm_mm']:.3f}",
                    f"{module3_results['n_tension_bolts']}",
                    f"{module3_results['T_total_kN']:.5f}",
                    f"{module3_results['T_per_bolt_kN']:.5f}",
                    f"{module3_results.get('tension_coord_mm', 0.0):.3f}",
                    str(module3_results["critical_bolts"]),
                ],
            })
            st.dataframe(mod3_df, use_container_width=True, hide_index=True)

            if not module3_results["tension_active"]:
                st.success(
                    "En este modelo preliminar uniaxial no se activa tracción en pernos."
                )
            else:
                st.warning(
                    "Se activa tracción en la fila extrema de pernos del lado que levanta. "
                    "El siguiente módulo verificará el espesor de placa."
                )
    
    with tab5:
        st.subheader("Módulo 4 - espesor mínimo de placa")

        if analysis_mode != "Uniaxial":
            st.info("En la barra lateral elegiste modo Biaxial. Cambia a Uniaxial para activar este módulo.")
        else:
            mod4_df = pd.DataFrame({
                "Parámetro": [
                    "B placa [mm]",
                    "N placa [mm]",
                    "bf columna [mm]",
                    "d columna [mm]",
                    "Voladizo mx [mm]",
                    "Voladizo my [mm]",
                    "Voladizo crítico m [mm]",
                    "Presión gobernante q_u [MPa]",
                    "Fy placa [MPa]",
                    "φ flexión",
                    "Espesor ingresado tp [mm]",
                    "Espesor requerido t_req [mm]",
                    "Utilización t_req/tp",
                    "Chequeo de espesor",
                ],
                "Valor": [
                    f"{module4_results['B_mm']:.3f}",
                    f"{module4_results['N_mm']:.3f}",
                    f"{module4_results['bf_mm']:.3f}",
                    f"{module4_results['d_mm']:.3f}",
                    f"{module4_results['mx_mm']:.3f}",
                    f"{module4_results['my_mm']:.3f}",
                    f"{module4_results['mcrit_mm']:.3f}",
                    f"{module4_results['q_u_MPa']:.5f}",
                    f"{module4_results['Fy_MPa']:.3f}",
                    f"{module4_results['phi_flexure']:.3f}",
                    f"{module4_results['tp_input_mm']:.3f}",
                    f"{module4_results['t_req_mm']:.3f}",
                    f"{module4_results['utilization']:.3f}",
                    "Cumple" if module4_results["thickness_ok"] else "No cumple",
                ],
            })
            st.dataframe(mod4_df, use_container_width=True, hide_index=True)

            if module4_results["thickness_ok"]:
                st.success(
                    "El espesor ingresado de la placa cumple con el espesor mínimo requerido en este modelo simplificado."
                )
            else:
                st.error(
                    "El espesor ingresado de la placa no cumple. Debe aumentarse el espesor o modificarse la geometría."
                )
    

    with tab6:
        st.subheader("Módulo 5 - acero del anclaje")

        if analysis_mode != "Uniaxial":
            st.info("En la barra lateral elegiste modo Biaxial. Cambia a Uniaxial para activar este módulo.")
        else:
            mod5_df = pd.DataFrame({
                "Parámetro": [
                    "Tipo de anclaje",
                    "Área efectiva tensión Ase,N [mm²]",
                    "Área efectiva cortante Ase,V [mm²]",
                    "fy anclaje [MPa]",
                    "fu anclaje ingresado [MPa]",
                    "futa efectivo usado [MPa]",
                    "φ acero tensión",
                    "φ acero cortante",
                    "Nsa nominal [kN]",
                    "Vsa nominal [kN]",
                    "φNsa [kN]",
                    "φVsa [kN]",
                    "Nua por perno [kN]",
                    "Vua por perno [kN]",
                    "Relación tensión",
                    "Relación cortante",
                    "Interacción acero",
                    "Chequeo tensión",
                    "Chequeo cortante",
                    "Chequeo interacción",
                    "Built-up grout pad",
                ],
                "Valor": [
                    module5_results["anchor_type"],
                    f"{module5_results['AseN_mm2']:.3f}",
                    f"{module5_results['AseV_mm2']:.3f}",
                    f"{module5_results['fya_MPa']:.3f}",
                    f"{module5_results['futa_input_MPa']:.3f}",
                    f"{module5_results['futa_eff_MPa']:.3f}",
                    f"{module5_results['phi_anchor_tension_steel']:.3f}",
                    f"{module5_results['phi_anchor_shear_steel']:.3f}",
                    f"{module5_results['Nsa_kN']:.5f}",
                    f"{module5_results['Vsa_kN']:.5f}",
                    f"{module5_results['phiNsa_kN']:.5f}",
                    f"{module5_results['phiVsa_kN']:.5f}",
                    f"{module5_results['Nua_per_bolt_kN']:.5f}",
                    f"{module5_results['Vua_per_bolt_kN']:.5f}",
                    f"{module5_results['tension_ratio']:.5f}",
                    f"{module5_results['shear_ratio']:.5f}",
                    f"{module5_results['interaction_value']:.5f}",
                    "Cumple" if module5_results["tension_ok"] else "No cumple",
                    "Cumple" if module5_results["shear_ok"] else "No cumple",
                    "Cumple" if module5_results["interaction_ok"] else "No cumple",
                    "Sí" if module5_results["use_built_up_grout_pad"] else "No",
                ],
            })
            st.dataframe(mod5_df, use_container_width=True, hide_index=True)

            if module5_results["interaction_ok"]:
                st.success("El acero del anclaje cumple preliminarmente en tensión, cortante e interacción.")
            else:
                st.error("El acero del anclaje no cumple preliminarmente. Revisa diámetro, área o cantidad de pernos.")

    with tab7:
        st.subheader("Módulo 6 - anclaje al concreto en tensión")

        if analysis_mode != "Uniaxial":
            st.info("En la barra lateral elegiste modo Biaxial. Cambia a Uniaxial para activar este módulo.")
        else:
            mod6_rows = [
                ["Tipo de anclaje", module6_results["anchor_type"]],
                ["Instalación", module6_results["anchor_installation"]],
                ["f'c [MPa]", f"{module6_results['fc_MPa']:.3f}"],
                ["hef [mm]", f"{module6_results['hef_mm']:.3f}"],
                ["da [mm]", f"{module6_results['da_mm']:.3f}"],
                ["kc", f"{module6_results['kc']:.3f}"],
                ["λa", f"{module6_results['lambda_a']:.3f}"],
                ["ψa", f"{module6_results['psi_a']:.3f}"],
                ["A_Nco [mm²]", f"{module6_results['ANco_mm2']:.3f}"],
                ["A_Nc [mm²]", f"{module6_results['ANc_mm2']:.3f}"],
                ["Edge OK en x (ca1 ≥ 1.5hef)", "Sí" if module6_results["edge_ok_x"] else "No"],
                ["Edge OK en y (ca1 ≥ 1.5hef)", "Sí" if module6_results["edge_ok_y"] else "No"],
                ["Supuesto sin borde crítico válido", "Sí" if module6_results["no_edge_effect_assumption_ok"] else "No"],
                ["Nb básico [kN]", f"{module6_results['Nb_kN']:.5f}"],
                ["Ncbg [kN]", f"{module6_results['Ncbg_kN']:.5f}"],
                ["φNcbg [kN]", f"{module6_results['phiNcbg_kN']:.5f}"],
                ["ψc,P", f"{module6_results['psi_c_P']:.3f}"],
                ["Npn [kN]", "-" if module6_results["Npn_kN"] is None else f"{module6_results['Npn_kN']:.5f}"],
                ["φNpn [kN]", "-" if module6_results["phiNpn_kN"] is None else f"{module6_results['phiNpn_kN']:.5f}"],
                ["Nsbg [kN]", "-" if module6_results["Nsbg_kN"] is None else f"{module6_results['Nsbg_kN']:.5f}"],
                ["φNsbg [kN]", "-" if module6_results["phiNsbg_kN"] is None else f"{module6_results['phiNsbg_kN']:.5f}"],
                ["Nua grupo [kN]", f"{module6_results['Nua_group_kN']:.5f}"],
                ["Resistencia gobernante φNn,cg [kN]", f"{module6_results['phiNn_cg_kN']:.5f}"],
                ["Chequeo concreto en tensión", "Cumple" if module6_results["concrete_tension_ok"] else "No cumple"],
            ]

            mod6_df = pd.DataFrame(mod6_rows, columns=["Parámetro", "Valor"])
            st.dataframe(mod6_df, use_container_width=True, hide_index=True)

            if not module6_results["no_edge_effect_assumption_ok"]:
                st.warning(
                    "Este Módulo 6 todavía está depurado para grupos sin borde crítico cercano. "
                    "Tus distancias a borde son menores que 1.5hef en al menos una dirección, así que "
                    "el resultado puede no ser representativo hasta que integremos los modificadores completos de borde."
                )

            if module6_results["anchor_type"] == "adhesive_anchor":
                st.warning(
                    "Para adhesive anchors, este módulo aún no calcula bond strength de 17.6.5. "
                    "Por ahora solo sirve como estructura de trabajo."
                )

            if module6_results["concrete_tension_ok"]:
                st.success("El anclaje al concreto cumple preliminarmente en tensión para los modos incluidos.")
            else:
                st.error("El anclaje al concreto no cumple preliminarmente en tensión para los modos incluidos.")

    with tab8:
        st.subheader("Módulo 7 - concreto en cortante")

        if analysis_mode != "Uniaxial":
            st.info("En la barra lateral elegiste modo Biaxial. Cambia a Uniaxial para activar este módulo.")
        else:
            mod7_rows = [
                ["φ concreto en cortante", f"{module7_results['phi_concrete_shear']:.3f}"],
                ["f'c [MPa]", f"{module7_results['fc_MPa']:.3f}"],
                ["hef [mm]", f"{module7_results['hef_mm']:.3f}"],
                ["λa", f"{module7_results['lambda_a']:.3f}"],
                ["ca1, shear [mm]", f"{module7_results['ca1_shear_mm']:.3f}"],
                ["Extensión perpendicular del grupo [mm]", f"{module7_results['group_span_perp_mm']:.3f}"],
                ["A_Vco [mm²]", f"{module7_results['AVco_mm2']:.3f}"],
                ["A_Vc [mm²]", f"{module7_results['AVc_mm2']:.3f}"],
                ["Vb básico [kN]", f"{module7_results['Vb_kN']:.5f}"],
                ["Vcbg [kN]", f"{module7_results['Vcbg_kN']:.5f}"],
                ["φVcbg [kN]", f"{module7_results['phiVcbg_kN']:.5f}"],
                ["kcp", f"{module7_results['kcp']:.3f}"],
                ["Vcpg [kN]", f"{module7_results['Vcpg_kN']:.5f}"],
                ["φVcpg [kN]", f"{module7_results['phiVcpg_kN']:.5f}"],
                ["Resistencia gobernante φVn,cg [kN]", f"{module7_results['phiVn_cg_kN']:.5f}"],
                ["Vua grupo [kN]", f"{module7_results['Vua_group_kN']:.5f}"],
                ["Chequeo concreto en cortante", "Cumple" if module7_results["shear_concrete_ok"] else "No cumple"],
                ["Nua grupo [kN]", f"{module7_results['Nua_group_kN']:.5f}"],
                ["φNn,cg [kN]", f"{module7_results['phiNn_cg_kN']:.5f}"],
                ["Interacción concreto N-V", f"{module7_results['interaction_concrete']:.5f}"],
                ["Chequeo interacción concreta", "Cumple" if module7_results["interaction_concrete_ok"] else "No cumple"],
                ["ha [mm]", f"{module7_results['member_thickness_for_shear_mm']:.3f}"],
                ["Advertencia de espesor de miembro", "Sí" if module7_results["thickness_warning"] else "No"],
            ]

            mod7_df = pd.DataFrame(mod7_rows, columns=["Parámetro", "Valor"])
            st.dataframe(mod7_df, use_container_width=True, hide_index=True)

            st.warning(
                "Este Módulo 7 usa una formulación simplificada provisional para concrete breakout en cortante, "
                "pensada para depuración del flujo del programa. Los modificadores completos de ACI 17.7.2 "
                "todavía no están integrados."
            )

            if module7_results["thickness_warning"]:
                st.warning(
                    "El espesor del miembro ha es menor que 1.5·ca1,shear. Esto puede afectar la validez "
                    "del shear breakout simplificado."
                )

            if module7_results["interaction_concrete_ok"]:
                st.success("El concreto cumple preliminarmente en cortante e interacción N-V.")
            else:
                st.error("El concreto no cumple preliminarmente en cortante o en interacción N-V.")

    with tab9:
        st.subheader("Módulo 8 - ACI 318-25 Sección 17.9")

        if analysis_mode != "Uniaxial":
            st.info("En la barra lateral elegiste modo Biaxial. Cambia a Uniaxial para activar este módulo.")
        else:
            mod8_rows = [
                ["Regla geométrica aplicada", module8_results["geometry_rule_label"]],
                ["Refuerzo suplementario para splitting", "Sí" if module8_results["supplementary_reinforcement_for_splitting"] else "No"],
                ["db [mm]", f"{module8_results['db_mm']:.3f}"],
                ["hef [mm]", f"{module8_results['hef_mm']:.3f}"],
                ["Espaciamiento real mínimo [mm]", f"{module8_results['min_spacing_real_mm']:.3f}"],
                ["Espaciamiento mínimo requerido [mm]", f"{module8_results['s_min_req_mm']:.3f}"],
                ["Chequeo de espaciamiento", "Cumple" if module8_results["spacing_ok"] else "No cumple"],
                ["Distancia real mínima al borde [mm]", f"{module8_results['min_edge_real_mm']:.3f}"],
                ["Distancia mínima requerida al borde [mm]", f"{module8_results['c_min_req_mm']:.3f}"],
                ["Chequeo de borde", "Cumple" if module8_results["edge_ok"] else "No cumple"],
                ["Recubrimiento requerido [mm]", f"{module8_results['required_cover_mm']:.3f}"],
                ["Tamaño máximo agregado [mm]", f"{module8_results['nominal_max_agg_mm']:.3f}"],
                ["Datos product-specific", "Sí" if module8_results["product_specific_geometry_data"] else "No"],
                ["hef límite 17.9.4 [mm]", "-" if module8_results["hef_limit_req_mm"] is None else f"{module8_results['hef_limit_req_mm']:.3f}"],
                ["Chequeo hef 17.9.4", "-" if module8_results["hef_limit_ok"] is None else ("Cumple" if module8_results["hef_limit_ok"] else "No cumple")],
                ["¿Se puede evaluar da' por 17.9.3?", "Sí" if module8_results["can_use_da_prime"] else "No"],
                ["da' estimado [mm]", "-" if module8_results["da_prime_mm"] is None else f"{module8_results['da_prime_mm']:.3f}"],
                ["Chequeo global geométrico", "Cumple" if module8_results["geometric_ok"] else "No cumple"],
            ]

            mod8_df = pd.DataFrame(mod8_rows, columns=["Parámetro", "Valor"])
            st.dataframe(mod8_df, use_container_width=True, hide_index=True)

            if module8_results["supplementary_reinforcement_for_splitting"]:
                st.info(
                    "ACI 17.9 permite salir de estos mínimos si existe refuerzo suplementario para controlar splitting. "
                    "Este módulo igual te muestra los mínimos de referencia."
                )

            if not module8_results["geometric_ok"]:
                st.warning(
                    "La geometría no cumple completamente con los mínimos de ACI 17.9. "
                    "Si el anclaje no produce splitting al instalarse y no será torqued, 17.9.3 permite usar un diámetro equivalente da'."
                )

            if module8_results["can_use_da_prime"] and module8_results["da_prime_mm"] is not None and module8_results["da_prime_mm"] < module8_results["db_mm"]:
                st.warning(
                    f"Para 17.9.3, el diámetro equivalente preliminar da' sería {module8_results['da_prime_mm']:.2f} mm, "
                    "y las fuerzas aplicadas deberían limitarse a las correspondientes a ese diámetro equivalente."
                )

            if module8_results["geometric_ok"]:
                st.success("El arreglo cumple los mínimos geométricos revisados de ACI 17.9 en esta etapa.")

    with tab10:
        st.subheader("Módulo 9 - transferencia de cortante en la base")

        if analysis_mode != "Uniaxial":
            st.info("En la barra lateral elegiste modo Biaxial. Cambia a Uniaxial para activar este módulo.")
        else:
            mod9_rows = [
                ["Eje uniaxial", module9_results["uniaxial_axis"]],
                ["Vu [kN]", f"{module9_results['Vu_kN']:.5f}"],
                ["Modo de evaluación", module9_results["base_shear_mode"]],
                ["Caso seleccionado", module9_results["selected_case"]],
                ["μ", f"{module9_results['mu']:.3f}"],
                ["φ fricción", f"{module9_results['phi_base_friction']:.3f}"],
                ["Pu en compresión [kN]", f"{module9_results['Pu_comp_kN']:.5f}"],
                ["φVn fricción [kN]", f"{module9_results['phiVn_friction_kN']:.5f}"],
                ["φVn grupo de anclajes [kN]", f"{module9_results['phiVn_anchor_group_kN']:.5f}"],
                ["φVn shear key/lug [kN]", f"{module9_results['phiVn_shear_key_kN']:.5f}"],
                ["φVn considerado [kN]", f"{module9_results['phiVn_selected_kN']:.5f}"],
                ["Mecanismos contribuyentes", ", ".join(module9_results["contributing_mechanisms"]) if module9_results["contributing_mechanisms"] else "-"],
                ["¿Se requiere shear key?", "Sí" if module9_results["shear_key_required"] else "No"],
                ["Vu remanente para shear key [kN]", f"{module9_results['Vu_remaining_for_key_kN']:.5f}"],
                ["Utilización Vu/φVn", f"{module9_results['utilization']:.5f}"],
                ["Chequeo del mecanismo actual", "Cumple" if module9_results["shear_ok"] else "No cumple"],
            ]

            mod9_df = pd.DataFrame(mod9_rows, columns=["Parámetro", "Valor"])
            st.dataframe(mod9_df, use_container_width=True, hide_index=True)

            if module9_results["mechanism_warning"] is not None:
                st.warning(module9_results["mechanism_warning"])

            st.info(
                "ACI 17.11 cubre las fallas del concreto de attachments con shear lugs, "
                "pero no cubre el diseño del acero ni de la soldadura de la placa y del lug."
            )

            if module9_results["shear_key_required"]:
                st.error(
                    f"Con la resistencia disponible sin shear key, no alcanza el cortante. "
                    f"La shear key debería resistir al menos {module9_results['Vu_remaining_for_key_kN']:.3f} kN."
                )
            else:
                st.success("No se requiere shear key en esta evaluación.")
    with tab11:
        st.subheader("Módulo 10 - soldadura")

        if analysis_mode != "Uniaxial":
            st.info("En la barra lateral elegiste modo Biaxial. Cambia a Uniaxial para activar este módulo.")
        else:
            st.markdown("### Soldadura columna - placa base")

            mod10_col_rows = [
                ["FEXX [MPa]", f"{module10_results['FEXX_MPa']:.3f}"],
                ["φ soldadura", f"{module10_results['phi_weld']:.3f}"],
                ["Demanda Tu columna-base [kN]", f"{module10_results['Tu_col_base_kN']:.5f}"],
                ["Tamaño filete [mm]", f"{module10_results['column_weld_size_mm']:.3f}"],
                ["Patrón solicitado", module10_results["column_weld_layout"]],
                ["Patrón usado", module10_results["column_weld_layout_used"]],
                ["Longitud geométrica [mm]", f"{module10_results['L_col_geom_mm']:.3f}"],
                ["β por longitud", f"{module10_results['beta_length']:.5f}"],
                ["l/w", f"{module10_results['l_over_w']:.3f}"],
                ["Longitud mínima 4w cumple", "Sí" if module10_results["min_length_ok"] else "No"],
                ["Reducción por longitud excesiva", "Sí" if module10_results["excessive_length_limit_applied"] else "No"],
                ["Longitud efectiva L_eff [mm]", f"{module10_results['L_col_weld_eff_mm']:.3f}"],
                ["Área efectiva Aw [mm²]", f"{module10_results['Aw_col_mm2']:.3f}"],
                ["φRn soldadura columna-base [kN]", f"{module10_results['phiRn_col_kN']:.5f}"],
                ["Utilización Tu/φRn", f"{module10_results['col_weld_util']:.5f}"],
                ["Chequeo columna-base", "Cumple" if module10_results["col_weld_ok"] else "No cumple"],
            ]

            st.dataframe(
                pd.DataFrame(mod10_col_rows, columns=["Parámetro", "Valor"]),
                use_container_width=True,
                hide_index=True,
            )
            
            if not module10_results["min_length_ok"]:
                st.warning(
                    "La longitud geométrica de soldadura es menor que 4 veces el tamaño del filete. "
                    "AISC J2.2b(c) exige ajustar el tamaño efectivo o aumentar la longitud."
                )

            if module10_results["excessive_length_limit_applied"]:
                st.warning(
                    "Se aplicó reducción de longitud efectiva por filete excesivamente largo conforme a J2-1 / límite de longitud."
                )

            if module10_results["col_weld_ok"]:
                st.success("La soldadura columna-placa cumple en este modelo preliminar de filete concéntrico.")
            else:
                st.error("La soldadura columna-placa no cumple en este modelo preliminar.")

            st.markdown("### Soldadura shear lug - placa base")

            mod10_lug_rows = [
                ["¿Se requiere shear key?", "Sí" if module10_results["shear_key_required"] else "No"],
                ["¿Se evaluó soldadura del lug?", "Sí" if module10_results["provide_shear_lug_weld"] else "No"],
                ["Demanda Vu lug [kN]", f"{module10_results['Vu_lug_kN']:.5f}"],
                ["Tamaño filete lug [mm]", f"{module10_results['shear_lug_weld_size_mm']:.3f}"],
                ["Longitud efectiva lug [mm]", f"{module10_results['shear_lug_weld_length_mm']:.3f}"],
                ["Área efectiva Aw lug [mm²]", f"{module10_results['Aw_lug_mm2']:.3f}"],
                ["φRn soldadura lug [kN]", f"{module10_results['phiRn_lug_kN']:.5f}"],
                ["Utilización Vu/φRn lug", f"{module10_results['lug_weld_util']:.5f}"],
                ["Chequeo lug-placa", "Cumple" if module10_results["lug_weld_ok"] else "No cumple"],
            ]

            st.dataframe(
                pd.DataFrame(mod10_lug_rows, columns=["Parámetro", "Valor"]),
                use_container_width=True,
                hide_index=True,
            )

            if module10_results["shear_key_required"] and not module10_results["provide_shear_lug_weld"]:
                st.warning(
                    "El Módulo 9 detectó que se requiere shear key, pero no activaste la revisión de soldadura shear lug-placa."
                )

            if module10_results["provide_shear_lug_weld"]:
                if module10_results["lug_weld_ok"]:
                    st.success("La soldadura shear lug-placa cumple en este modelo preliminar.")
                else:
                    st.error("La soldadura shear lug-placa no cumple en este modelo preliminar.")

            st.info(
                "Este módulo usa un modelo preliminar de filete concéntricamente cargado basado en AISC J2. "
                "Para grupos excéntricos, el Manual AISC permite usar métodos de grupo de soldadura/centro instantáneo."
            )

            st.warning(
                "Si la base forma parte de un sistema sísmico especial, las soldaduras de column-to-base plate "
                "pueden tener requisitos adicionales en AISC 341 / AWS D1.8."
            )

    with tab12:
        st.subheader("Módulo 11 - análisis biaxial preliminar")

        if analysis_mode != "Biaxial":
            st.info("En la barra lateral elegiste modo Uniaxial. Cambia a Biaxial para activar este módulo.")
        else:
            mod11_rows = [
                ["Pu [kN]", f"{module11_results['Pu_kN']:.5f}"],
                ["Mux [kN·m]", f"{module11_results['Mux_kNm']:.5f}"],
                ["Muy [kN·m]", f"{module11_results['Muy_kNm']:.5f}"],
                ["Área de placa [mm²]", f"{module11_results['A_mm2']:.3f}"],
                ["Ix [mm⁴]", f"{module11_results['Ix_mm4']:.3f}"],
                ["Iy [mm⁴]", f"{module11_results['Iy_mm4']:.3f}"],
                ["e_x = Muy/Pu [mm]", f"{module11_results['e_x_mm']:.3f}"],
                ["e_y = Mux/Pu [mm]", f"{module11_results['e_y_mm']:.3f}"],
                ["q_max [MPa]", f"{module11_results['q_max_MPa']:.5f}"],
                ["q_min [MPa]", f"{module11_results['q_min_MPa']:.5f}"],
                ["Compresión total preliminar", "Sí" if module11_results["full_compression"] else "No"],
                ["Posible levantamiento preliminar", "Sí" if module11_results["possible_uplift"] else "No"],
                ["Tracción total preliminar [kN]", f"{module11_results['T_total_prelim_kN']:.5f}"],
                ["Perno crítico preliminar", "-" if module11_results["critical_bolt"] is None else str(module11_results["critical_bolt"])],
                ["Tracción máxima preliminar por perno [kN]", f"{module11_results['Tmax_prelim_kN']:.5f}"],
            ]

            st.dataframe(
                pd.DataFrame(mod11_rows, columns=["Parámetro", "Valor"]),
                use_container_width=True,
                hide_index=True,
            )

            st.markdown("### Presiones en las esquinas")
            corner_df = pd.DataFrame(
                [
                    {"Esquina": k, "q [MPa]": v}
                    for k, v in module11_results["corner_pressures_MPa"].items()
                ]
            )
            st.dataframe(corner_df, use_container_width=True, hide_index=True)

            st.markdown("### Tracción preliminar en pernos")
            st.dataframe(
                module11_results["bolt_tension_df"][["Perno", "x_mm", "y_mm", "demand_index", "T_prelim_kN"]],
                use_container_width=True,
                hide_index=True,
            )

            if module11_results["possible_uplift"]:
                st.warning(
                    "Este módulo detecta posible levantamiento en biaxial y asigna una tracción preliminar elástica a los pernos. "
                    "Es una aproximación útil para depuración, no el equilibrio final no lineal."
                )
            else:
                st.success("En esta evaluación preliminar biaxial, toda la placa permanece en compresión.")
   
    with tab13:
        st.subheader("Módulo 12 - refinamiento biaxial por malla")

        if analysis_mode != "Biaxial":
            st.info("En la barra lateral elegiste modo Uniaxial. Cambia a Biaxial para activar este módulo.")
        else:
            mod12_rows = [
                ["nx", f"{module12_results['nx']}"],
                ["ny", f"{module12_results['ny']}"],
                ["dx [mm]", f"{module12_results['dx_mm']:.3f}"],
                ["dy [mm]", f"{module12_results['dy_mm']:.3f}"],
                ["dA [mm²]", f"{module12_results['dA_mm2']:.3f}"],
                ["q_max [MPa]", f"{module12_results['q_max_MPa']:.5f}"],
                ["q_min [MPa]", f"{module12_results['q_min_MPa']:.5f}"],
                ["Compresión total", "Sí" if module12_results["full_compression"] else "No"],
                ["Posible levantamiento", "Sí" if module12_results["possible_uplift"] else "No"],
                ["C total [kN]", f"{module12_results['C_kN']:.5f}"],
                ["T equivalente [kN]", f"{module12_results['T_equiv_kN']:.5f}"],
                ["xC [mm]", f"{module12_results['xC_mm']:.5f}"],
                ["yC [mm]", f"{module12_results['yC_mm']:.5f}"],
                ["Mx interno [kN·m]", f"{module12_results['Mx_internal_kNm']:.5f}"],
                ["My interno [kN·m]", f"{module12_results['My_internal_kNm']:.5f}"],
                ["Residual Mx [kN·m]", f"{module12_results['Mx_residual_kNm']:.5f}"],
                ["Residual My [kN·m]", f"{module12_results['My_residual_kNm']:.5f}"],
                ["Perno crítico refinado", "-" if module12_results["critical_bolt"] is None else str(module12_results["critical_bolt"])],
                ["Tracción máxima refinada por perno [kN]", f"{module12_results['Tmax_refined_kN']:.5f}"],
            ]

            st.dataframe(
                pd.DataFrame(mod12_rows, columns=["Parámetro", "Valor"]),
                use_container_width=True,
                hide_index=True,
            )

            st.markdown("### Tracción refinada en pernos")
            st.dataframe(
                module12_results["bolt_tension_df"][["Perno", "x_mm", "y_mm", "uplift_index", "T_refined_kN"]],
                use_container_width=True,
                hide_index=True,
            )

            st.warning(
                "Este módulo mejora la rama biaxial usando integración por malla y separación entre zonas comprimidas y levantadas. "
                "Aun así, no representa todavía la solución final exacta con equilibrio no lineal completo."
            )   
    with tab14:
        st.subheader("Módulo 13 - cierre del diseño")

        st.markdown("### Resumen global de chequeos")
        st.dataframe(
            module13_results["summary_df"],
            use_container_width=True,
            hide_index=True,
        )

        st.markdown("### Estado global")

        if module13_results["global_status"] == "Cumple":
            st.success("Estado global: Cumple")
        elif module13_results["global_status"] == "Preliminarmente aceptable":
            st.info("Estado global: Preliminarmente aceptable")
        elif module13_results["global_status"] == "Revisión requerida":
            st.warning("Estado global: Revisión requerida")
        elif module13_results["global_status"] == "No cumple":
            st.error("Estado global: No cumple")
        else:
            st.info(f"Estado global: {module13_results['global_status']}")

        st.markdown("### Chequeos críticos")
        if module13_results["critical_checks"]:
            critical_df = pd.DataFrame(module13_results["critical_checks"])
            st.dataframe(critical_df, use_container_width=True, hide_index=True)
        else:
            st.success("No se identificaron chequeos críticos.")    
    with tab15:
        st.subheader("Módulo 14 - gráficas de presión y tracción")

        if analysis_mode == "Uniaxial":
            st.markdown("### Distribución de presión uniaxial")
            st.pyplot(module14_figs["uniaxial_pressure"], clear_figure=False)

            st.markdown("### Tracción en pernos")
            st.pyplot(module14_figs["uniaxial_anchor_tension"], clear_figure=False)

        elif analysis_mode == "Biaxial":
            if "biaxial_elastic_pressure" in module14_figs:
                st.markdown("### Presión biaxial elástica")
                st.pyplot(module14_figs["biaxial_elastic_pressure"], clear_figure=False)

            if "biaxial_contact_pressure" in module14_figs:
                st.markdown("### Presión de contacto biaxial")
                st.pyplot(module14_figs["biaxial_contact_pressure"], clear_figure=False)

            st.markdown("### Tracción refinada en pernos")
            st.pyplot(module14_figs["biaxial_anchor_tension"], clear_figure=False)
    with tab16:
        st.subheader("Resumen")

        st.write("Esta pestaña permite generar la memoria de cálculo en formato Word.")

        if st.button("Generar memoria de cálculo en Word"):
            report_path = module15_generate_word_report(
                analysis_mode=analysis_mode,
                loads=loads,
                materials=materials,
                column_plot=column_plot,
                base_plate=base_plate,
                anchors=anchors,
                pedestal=pedestal,
                module13_results=module13_results,
                module14_figs=module14_figs,
                module2_results=module2_results if analysis_mode == "Uniaxial" else None,
                module4_results=module4_results if analysis_mode == "Uniaxial" else None,
                module5_results=module5_results if analysis_mode == "Uniaxial" else None,
                module6_results=module6_results if analysis_mode == "Uniaxial" else None,
                module7_results=module7_results if analysis_mode == "Uniaxial" else None,
                module8_results=module8_results if analysis_mode == "Uniaxial" else None,
                module9_results=module9_results if analysis_mode == "Uniaxial" else None,
                module10_results=module10_results if analysis_mode == "Uniaxial" else None,
                module11_results=module11_results if analysis_mode == "Biaxial" else None,
                module12_results=module12_results if analysis_mode == "Biaxial" else None,
            )

            st.success("Memoria generada correctamente.")

            with open(report_path, "rb") as f:
                st.download_button(
                    label="Descargar memoria .docx",
                    data=f,
                    file_name="memoria_placa_base.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
    with tab17:
        st.subheader("Estado del desarrollo")
        st.write("**Módulos activos:**")
        st.write("- Rama uniaxial: Módulos 1 a 10")
        st.write("- Rama biaxial: Módulos 11 y 12")
        st.write("- Cierre global: Módulo 13")
        st.write("- Gráficas automáticas: Módulo 14")
        st.write("- Memoria en Word: Módulo 15")
        st.write("**Siguientes mejoras recomendadas:**")
        st.write("- Ecuaciones OMML nativas de Word")
        st.write("- Exportación PDF")
        st.write("- Refinamiento no lineal del biaxial")

except Exception as exc:
    st.error(f"Error en los datos de entrada: {exc}")